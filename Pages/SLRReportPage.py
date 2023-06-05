import ast
import numbers
import os
import re
import time
import glob

import docx
import openpyxl
import pandas as pd
from pathlib import Path
from re import search
from selenium.common import NoSuchElementException
from selenium.webdriver.common.by import By
from selenium.webdriver.support import expected_conditions as ec
from selenium.webdriver.support.wait import WebDriverWait
from selenium.webdriver.support.ui import Select
from selenium.webdriver import ActionChains
from pandas.core.common import flatten

from Pages.Base import Base, fWaitFor
from Pages.ExtendedBasePage import ExtendedBase
from Pages.ImportPublicationsPage import ImportPublicationPage
from Pages.OpenLiveSLRPage import LiveSLRPage
from utilities.customLogger import LogGen
from utilities.logScreenshot import cLogScreenshot
from utilities.readProperties import ReadConfig


class SLRReport(Base):

    """Constructor of the SLRReport class"""
    def __init__(self, driver, extra):
        # initializing the driver from base class
        super().__init__(driver, extra)  
        self.extra = extra
        # Instantiate the Base class
        self.base = Base(self.driver, self.extra)
        # Creating object of liveslrpage class
        self.liveslrpage = LiveSLRPage(self.driver, self.extra)
        # Instantiate the logger class
        self.logger = LogGen.loggen()
        # Instantiate the logScreenshot class
        self.LogScreenshot = cLogScreenshot(self.driver, self.extra)
        # Creating object of ExtendedBase class
        self.exbase = ExtendedBase(self.driver, extra)  
        # Creating object of ImportPublicationPage class
        self.imppubpage = ImportPublicationPage(self.driver, extra)                     
        # Instantiate webdriver wait class
        self.wait = WebDriverWait(driver, 20)

    def select_data(self, locator, locator_button, env):
        if self.isselected(locator_button, env):
            self.LogScreenshot.fLogScreenshot(message=f"Selected Element: {locator}",
                                              pass_=True, log=True, screenshot=True)
        else:
            self.jsclick(locator, env, UnivWaitFor=10)
            if self.isselected(locator_button, env):
                self.LogScreenshot.fLogScreenshot(message=f"Selected Element: {locator}",
                                                  pass_=True, log=True, screenshot=True)

    def select_sub_section(self, locator, locator_button, env, scroll=None):
        if self.scroll(scroll, env, UnivWaitFor=60):
            if self.isselected(locator_button, env):
                self.LogScreenshot.fLogScreenshot(message=f"{locator} already selected",
                                                  pass_=True, log=True, screenshot=True)
            else:
                self.jsclick(locator, env, UnivWaitFor=10)
                if self.isselected(locator_button, env):
                    self.LogScreenshot.fLogScreenshot(message=f"{locator} selected",
                                                      pass_=True, log=True, screenshot=True)
            self.scrollback("SLR_page_header", env)

    def select_all_sub_section(self, locator, locator_button, env, scroll=None):
        if self.scroll(scroll, env, UnivWaitFor=20):
            if self.isselected(locator_button, env):
                self.LogScreenshot.fLogScreenshot(message=f"{locator} already selected",
                                                  pass_=True, log=True, screenshot=True)
            else:
                self.jsclick(locator, env, UnivWaitFor=10)
                if self.isselected(locator_button, env):
                    self.LogScreenshot.fLogScreenshot(message=f"{locator} selected",
                                                      pass_=True, log=True, screenshot=True)
            self.scrollback("SLR_page_header", env)

    def get_additional_criteria_values(self, locator_study, locator_var, env):
        ele1 = self.select_elements(locator_study, env)
        ele2 = self.select_elements(locator_var, env)
        return ele1, ele2

    def get_source_template(self, filepath, col_name):
        file = pd.read_excel(filepath)
        expectedfilepath = list(os.getcwd()+file[col_name].dropna())
        return expectedfilepath
    
    def get_duplicates_from_list(self, input_list):
        return list(set([x for x in input_list if input_list.count(x) > 1]))

    def list_comparison_between_reports_data(self, source_list, compex_list, webex_list=None, word=None):
        idx = 0
        res_index = []
        res_list = []
        if webex_list is not None and word is not None:
            for i in source_list:
                if i != compex_list[idx] and i != webex_list[idx] and i != word[idx]:
                    res_index.append(idx)
                idx = idx + 1
        elif webex_list is None and word is not None:
            for i in source_list:
                if i != compex_list[idx] and i != word[idx]:
                    res_index.append(idx)
                idx = idx + 1
        elif webex_list is not None and word is None:
            for i in source_list:
                if i != compex_list[idx] and i != webex_list[idx]:
                    res_index.append(idx)
                idx = idx + 1
        else:
            for i in source_list:
                if i != compex_list[idx]:
                    res_index.append(idx)
                idx = idx + 1

        if webex_list is not None and word is not None:
            for index, n in enumerate(res_index):
                res_list.append([source_list[n]])
                res_list[index].append(compex_list[n])
                res_list[index].append(webex_list[n])
                res_list[index].append(word[n])
        elif webex_list is None and word is not None:
            for index, n in enumerate(res_index):
                res_list.append([source_list[n]])
                res_list[index].append(compex_list[n])
                res_list[index].append(word[n])
        elif webex_list is not None and word is None:
            for index, n in enumerate(res_index):
                res_list.append([source_list[n]])
                res_list[index].append(compex_list[n])
                res_list[index].append(webex_list[n])
        else:
            for index, n in enumerate(res_index):
                res_list.append([source_list[n]])
                res_list[index].append(compex_list[n])

        return res_list
    
    def validate_additional_criteria_val(self, filepath, locator_study, locator_var, env):
        # Read reportedvariables and studydesign expected data values
        design_val, var_val = self.liveslrpage.get_data_values(filepath)
        # Get the actual values
        act_study_design, act_rep_var = self.get_additional_criteria_values(locator_study, locator_var, env)
        try:
            for k in act_study_design:
                if k.text in design_val:
                    self.LogScreenshot.fLogScreenshot(
                        message=f"Correct StudyDesign Value is displayed in NMA Page:{k.text}",
                        pass_=True, log=True, screenshot=False)
            for v in act_rep_var:
                if v.text in var_val:
                    self.LogScreenshot.fLogScreenshot(
                        message=f"Correct Reported Variable is displayed in NMA Page:{v.text}",
                        pass_=True, log=True, screenshot=False)
        except Exception:
            raise Exception("Unable to validate Additional Criteria Values")

    def validate_selected_area(self, pop, slr, env):
        pop_sel, slr_sel = self.collect_selected_area_details("selected_area_population",
                                                              "selected_area_slrtype", env)
        if pop in pop_sel and slr in slr_sel:
            self.LogScreenshot.fLogScreenshot(message=f"Selected elements are shown correct\n"
                                                      f"Selected Area value: {pop_sel} {slr_sel}\n"
                                                      f"Selected Element value: {pop}, {slr}",
                                              pass_=True, log=True, screenshot=False)
        else:
            self.LogScreenshot.fLogScreenshot(message=f"Selected elements are wrong",
                                              pass_=False, log=True, screenshot=False)
            raise Exception("Selected Population and SLR Type are not matching in Selected area")

    def prism_value_validation(self, prism, efilename, wfilename, word_filename, env):

        # Reading Study Identifier column values from WebExcel Sheet
        webexcel = openpyxl.load_workbook(f'ActualOutputs//{wfilename}')
        webcount = []
        webcount_final = []

        for sheet in webexcel.sheetnames:
            webexcel_value = pd.read_excel(f'ActualOutputs//{wfilename}', sheet_name=sheet, skiprows=3)
            webex = webexcel_value['Study Identifier']
            webcount.append([item for item in webex if str(item) != 'nan'])

        # Removing duplicates to get the proper length of Study Identifier data
        [webcount_final.append(x) for x in list(flatten(webcount)) if x not in webcount_final]

        # Read PRISMA value from Complete Excel Sheet
        excel = openpyxl.load_workbook(f'ActualOutputs//{efilename}')
        excel_sheet = excel['Updated PRISMA']
        count_str = excel_sheet['B22'].value
        count_val = excel_sheet['C22'].value

        # Reading PRISMA Value from Complete Word Report
        docs = docx.Document(f'ActualOutputs//{word_filename}')
        word = []
        for row in docs.tables[4].rows:
            word.append(row.cells[1].text)

        self.LogScreenshot.fLogScreenshot(message=f"WebExcel FileName is: {wfilename} and Count Value is: "
                                                  f"{len(webcount_final)} \nExcel FileName is: {efilename} and "
                                                  f"Count value is: {count_str} {count_val}"
                                                  f"Word FileName is: {word_filename}\n and Count value is: {word[8]}",
                                          pass_=True, log=True, screenshot=False)

        if int(prism) == count_val and len(webcount_final) == count_val and int(word[8]) == count_val:
            self.LogScreenshot.fLogScreenshot(message=f"WebExcel Prisma Count Value: {len(webcount_final)}\n"
                                                      f"Excel Sheet Prisma Count Value: {count_val}\n"
                                                      f"Word Prisma Count Value: {word[8]}\n"
                                                      f"UI Updated Prisma Count Value: {prism}\n"
                                                      f"Records are matching",
                                              pass_=True, log=True, screenshot=True)
        else:
            self.LogScreenshot.fLogScreenshot(message=f"WebExcel Prisma Count Value: {len(webcount_final)}\n"
                                                      f"Excel Sheet Prisma Count Value: {count_val}\n"
                                                      f"Word Prisma Count Value: {word[8]}\n"
                                                      f"UI Updated Prisma Count Value: {prism}\n"
                                                      f"Records are not matching",
                                              pass_=False, log=True, screenshot=False)
            raise Exception("Prisma count values are mismatching")
        self.scrollback("SLR_page_header", env)

    def collect_selected_area_details(self, locator1, locator2, env):
        x = self.get_text(locator1, env)
        y = self.get_text(locator2, env)
        return x, y

    def preview_result(self, locator, env):
        if self.clickable(locator, env):
            self.jsclick(locator, env)
            self.LogScreenshot.fLogScreenshot(message=f"{locator} is clickable",
                                              pass_=True, log=True, screenshot=False)
        else:
            self.LogScreenshot.fLogScreenshot(message=f"{locator} is not clickable",
                                              pass_=False, log=True, screenshot=False)

    def table_display_check(self, locator, env):
        if self.isdisplayed(locator, env, UnivWaitFor=120):
            self.LogScreenshot.fLogScreenshot(message=f"{locator} is displayed",
                                              pass_=True, log=True, screenshot=True)
        else:
            time.sleep(10)
            self.driver.find_element(getattr(By, self.locatortype(locator, env)), self.locatorpath(locator, env))\
                .is_displayed()
            self.LogScreenshot.fLogScreenshot(message=f"{locator} is displayed with extra wait time",
                                              pass_=True, log=True, screenshot=True)

    # def check_download_error_ele(self, locator):
    #     return self.wait.until(ec.invisibility_of_element_located((By.XPATH, locator)))
    #     # if self.isdisplayed(locator, UnivWaitFor=5):
    #     #     self.LogScreenshot.fLogScreenshot(message=f"{locator} is displayed",
    #     #                                       pass_=False, log=True, screenshot=True)
    #     #     return True
    #     # else:
    #     #     return False

    @fWaitFor
    def check_download_error_element(self, locator, UnivWaitFor=0):
        try:
            if self.isdisplayed(locator):
                # value = True
                value = self.wait.until(
                    ec.presence_of_element_located((getattr(By, self.locatortype(locator)), self.locatorpath(locator))))
                self.LogScreenshot.fLogScreenshot(message=f"Entered into IF block",
                                                  pass_=False, log=True, screenshot=False)
                self.LogScreenshot.fLogScreenshot(message=f"IF block value is: {value}",
                                                  pass_=False, log=True, screenshot=True)
            else:
                # time.sleep(5)
                value = self.driver.find_element(getattr(By, self.locatortype(locator)),
                                                 self.locatorpath(locator)).is_displayed()
                self.LogScreenshot.fLogScreenshot(message=f"Entered into ELSE block",
                                                  pass_=False, log=True, screenshot=False)
                self.LogScreenshot.fLogScreenshot(message=f"ELSE block value is: {value}",
                                                  pass_=False, log=True, screenshot=True)
            return value
        except NoSuchElementException:
            self.LogScreenshot.fLogScreenshot(message=f"{locator} is not present",
                                              pass_=False, log=True, screenshot=False)
            return False

    @fWaitFor
    def generate_download_report(self, locator, env, UnivWaitFor=0):
        # download table csv
        try:
            # Get list of files before downloading the new report
            list_of_files_before_download = glob.glob('ActualOutputs//*')

            # Click on downlaod button
            self.jsclick(locator, env)
            # wait to download the report
            time.sleep(15)
            # Get list of files after downloading the new report
            list_of_files_after_download = glob.glob('ActualOutputs//*')

            if len(list_of_files_after_download) > len(list_of_files_before_download):
                self.LogScreenshot.fLogScreenshot(message=f"Report download is success",
                                                  pass_=True, log=True, screenshot=False)
            else:
                time.sleep(15)
                # Get list of files after downloading the new report with extra wait time
                list_of_files_after_download = glob.glob('ActualOutputs//*')
                if len(list_of_files_after_download) > len(list_of_files_before_download):
                    self.LogScreenshot.fLogScreenshot(message=f"Report download is success with extra wait time",
                                                      pass_=True, log=True, screenshot=False)
                else:
                    self.LogScreenshot.fLogScreenshot(message=f"Report is not downloaded",
                                                      pass_=False, log=True, screenshot=False)
                    raise Exception("Report download is failed")                                                                                        
        except Exception:
            self.LogScreenshot.fLogScreenshot(message=f"Error in downloading the table",
                                              pass_=False, log=True, screenshot=True)
            raise Exception("Download table failed")

    def back_to_report_page(self, locator, env):
        self.jsclick(locator, env)

    @fWaitFor
    def get_latest_filename(self, UnivWaitFor=0):
        list_of_files = glob.glob('ActualOutputs//*')
        # Get the latest downloaded file name with full path based on the downloaded time
        latest_file_path = max(list_of_files, key=os.path.getctime)
        # Extracting the filename from the latest file path
        latest_filename = os.path.basename(latest_file_path)
        # self.LogScreenshot.fLogScreenshot(message=f"Latest Filename from Actual Outputs folder is : "
        #                                           f"{latest_filename}", pass_=True, log=True, screenshot=False)
        return latest_filename                                                
    
    def get_and_validate_filename(self, filepath):
        try:
            filename = self.get_latest_filename(UnivWaitFor=180)
            file = pd.read_excel(filepath)
            expectedname = list(file['ExpectedFilenames'].dropna())
            # actualname = [x for x in re.compile(r'\D+').findall(filename)]
            actualname = filename[:-19]
            # for i in expectedname:
            #     # if i == actualname[0]:
            #     if i == actualname:
            #         self.LogScreenshot.fLogScreenshot(message=f"Correct file is downloaded",
            #                                           pass_=True, log=True, screenshot=False)
            #         break
            if actualname in expectedname:
                self.LogScreenshot.fLogScreenshot(message=f"Correct file is downloaded. Filename is {filename}",
                                                  pass_=True, log=True, screenshot=False)
                return filename
            else:
                raise Exception
        except Exception:
            self.LogScreenshot.fLogScreenshot(message=f"Filename is not present in the expected list. Expected "
                                                      f"Filenames are {expectedname} and Actual "
                                                      f"Filename is {actualname}",
                                                      pass_=False, log=True, screenshot=False)
            raise Exception("Error during filename validation")

    # method to get the downloaded file name
    def getFilenameAndValidate(self, waitTime):
        self.driver.execute_script("window.open()")
        # switch to new tab
        self.driver.switch_to.window(self.driver.window_handles[-1])
        # navigate to chrome downloads
        self.driver.get('chrome://downloads')
        # define the endTime
        endTime = time.time() + waitTime
        while True:
            try:
                time.sleep(10)
                filename = self.driver.execute_script(
                    "return document.querySelector('downloads-manager').shadowRoot.querySelector"
                    "('#downloadsList downloads-item').shadowRoot.querySelector('div#content  #file-link').text")
                self.LogScreenshot.fLogScreenshot(message=f"Downloaded filename is {filename}",
                                                  pass_=True, log=True, screenshot=False)
                self.driver.close()
                self.driver.switch_to.window(self.driver.window_handles[1])
                return filename
            except Exception:
                pass
            time.sleep(1)
            if time.time() > endTime:
                break
    
    def excel_content_validation(self, source_template, index, webexcel_filename, excel_filename, first_col):
        self.LogScreenshot.fLogScreenshot(message=f"Content validation between Expected Excel File and downloaded "
                                                  f"WebExcel and Complete Excel Reports",
                                          pass_=True, log=True, screenshot=False)

        self.LogScreenshot.fLogScreenshot(message=f"Expected Excel Filename is: {Path(f'{source_template[0]}').stem}, "
                                                  f"Downloaded FileNames are: {webexcel_filename} "
                                                  f"and \n{excel_filename}",
                                          pass_=True, log=True, screenshot=False)

        '''Reference-LIVEHTA1161: Check position of 'Back to TOC' button and presence of column names at row 4'''
        wb = openpyxl.load_workbook(f'ActualOutputs//{webexcel_filename}')
        for i in wb.sheetnames:
            # Check the position of 'Back to TOC' button
            ex = openpyxl.load_workbook(f'ActualOutputs//{excel_filename}')
            ex_sheet = ex[i]
            if ex_sheet['H1'].value == 'Back To Toc':
                self.LogScreenshot.fLogScreenshot(message=f"'Back To Toc' option is present in '{i}' sheet",
                                                  pass_=True, log=True, screenshot=False)
            else:
                self.LogScreenshot.fLogScreenshot(message=f"'Back To Toc' option is not present in '{i}' "
                                                          f"sheet",
                                                  pass_=False, log=True, screenshot=False)
                raise Exception(f"'Back To Toc' option is not present in '{i}' sheet")
            
            # Check the presence of column names at row 4
            df = pd.read_excel(f'ActualOutputs//{excel_filename}', sheet_name=i, skiprows=3)
            if first_col in df.columns.values:
                self.LogScreenshot.fLogScreenshot(message=f"Column names are present at row 4.",
                                                  pass_=True, log=True, screenshot=False)
            else:
                self.LogScreenshot.fLogScreenshot(message=f"Column names are not present at row 4",
                                                  pass_=False, log=True, screenshot=False)
                raise Exception(f"Column names are not present at row 4")           

        # Actual excel content validation step starts here
        source_data = openpyxl.load_workbook(f'{source_template[0]}')

        self.LogScreenshot.fLogScreenshot(message=f"Expected Excel file Sheetnames are: {source_data.sheetnames}",
                                          pass_=True, log=True, screenshot=False)
        
        expected_data = pd.read_excel(f'{source_template[0]}', sheet_name=source_data.sheetnames[index])
        webexcel = pd.concat(pd.read_excel(f'ActualOutputs//{webexcel_filename}', sheet_name=None, skiprows=3),
                             ignore_index=True)
        excel = pd.concat(pd.read_excel(f'ActualOutputs//{excel_filename}', sheet_name=None, skiprows=3),
                          ignore_index=True)

        try:
            # Check the length of 1st column from the report to make sure number of rows are as expected
            source_len = expected_data[first_col]
            webex_len = webexcel[first_col]
            compex_len = excel[first_col]

            source_len = [item for item in source_len if str(item) != 'nan']
            webex_len = [item for item in webex_len if str(item) != 'nan']
            compex_len = [item for item in compex_len if str(item) != 'nan']

            if len(source_len) == len(webex_len) == len(compex_len):
                self.LogScreenshot.fLogScreenshot(message=f"Elements length is matching between Expected Excel, "
                                                          f"Web_Excel and Complete Excel Report. "
                                                          f"Expected Excel Elements Length: {len(source_len)}\n "
                                                          f"WebExcel Elements Length: {len(webex_len)}\n "
                                                          f"Excel Elements Length: {len(compex_len)}\n",
                                                  pass_=True, log=True, screenshot=False)

                # Content validation starts from here
                for col in expected_data.columns.values:
                    source = expected_data[col]
                    webex = webexcel[col]
                    compex = excel[col]

                    source = [item for item in source if str(item) != 'nan']
                    webex = [item for item in webex if str(item) != 'nan']
                    compex = [item for item in compex if str(item) != 'nan']

                    comparison_result = self.list_comparison_between_reports_data(source, compex, webex_list=webex)

                    if len(comparison_result) == 0:
                        self.LogScreenshot.fLogScreenshot(message=f"From '{source_data.sheetnames[index]}' Report, "
                                                                  f"Values in Column '{col}' are matching between "
                                                                  f"Expected Excel, WebExcel and Complete Excel "
                                                                  f"Report.\n",
                                                          pass_=True, log=True, screenshot=False)
                    else:
                        self.LogScreenshot.fLogScreenshot(message=f"From '{source_data.sheetnames[index]}' Report, "
                                                                  f"Values in Column '{col}' are not matching between "
                                                                  f"Expected Excel, WebExcel and Complete Excel "
                                                                  f"Report. Mismatch values are arranged in following "
                                                                  f"order -> Expected Excel File, WebExcel, Complete "
                                                                  f"Excel. {comparison_result}",
                                                          pass_=False, log=True, screenshot=False)
                        raise Exception("Elements are not matching between Webexcel and Complete Excel")
            else:
                self.LogScreenshot.fLogScreenshot(message=f"Elements length is not matching between Expected Excel, "
                                                          f"Web_Excel and Complete Excel Report. "
                                                          f"Expected Excel Elements Length: {len(source_len)}\n "
                                                          f"WebExcel Elements Length: {len(webex_len)}\n "
                                                          f"Excel Elements Length: {len(compex_len)}\n",
                                                  pass_=False, log=True, screenshot=False)
                raise Exception(f"Elements length is not matching between Expected Excel, Web_Excel and "
                                f"Complete Excel Report.")
        except Exception:
            raise Exception("Error in Excel sheet content validation")

    def word_content_validation(self, filepath, slr_type_index, word_filename):
        self.LogScreenshot.fLogScreenshot(message=f"Content validation between Expected Excel File and "
                                                  f"Complete Word Report",
                                          pass_=True, log=True, screenshot=False)

        source_template = self.get_source_template(filepath, 'ExpectedSourceTemplateFile_Word')

        self.LogScreenshot.fLogScreenshot(message=f"FileName is: {word_filename}",
                                          pass_=True, log=True, screenshot=False)

        source_excel = openpyxl.load_workbook(f'{source_template[slr_type_index]}')                                          
        
        # Index of Table number 6 is : 5. Starting point for word table content comparison
        table_count = 5

        self.LogScreenshot.fLogScreenshot(message=f"Expected Excel file Sheetnames are: {source_excel.sheetnames}",
                                          pass_=True, log=True, screenshot=False)
        for sheet in source_excel.sheetnames:
            src_data = pd.read_excel(f'{source_template[slr_type_index]}', sheet_name=sheet)
            docs = docx.Document(f'ActualOutputs//{word_filename}')
            try:
                table = docs.tables[table_count]
                data = [[cell.text for cell in row.cells] for row in table.rows]
                df_word = pd.DataFrame(data)

                # Check the length of 1st column from the report to make sure number of rows are as expected
                src_data_len = src_data[df_word.values[0][0]]
                word_len = []
                for row in docs.tables[table_count].rows:
                    word_len.append(row.cells[0].text)

                src_data_len = [item for item in src_data_len if str(item) != 'nan']
                word_len.pop(0)

                if len(src_data_len) == len(word_len):
                    self.LogScreenshot.fLogScreenshot(message=f"Elements length is matching between Expected Excel and "
                                                              f"Complete Word Report. Expected Excel Elements Length: "
                                                              f"{len(src_data_len)}\n Word Elements Length: "
                                                              f"{len(word_len)}\n",
                                                      pass_=True, log=True, screenshot=False)

                    # Content validation starts from here
                    # Using count variable to loop over columns in word document
                    count = 0
                    # df_word.values[0] will give the list of column names from the table in Word document
                    for col_name in df_word.values[0]:
                        # # Restricting comparison upto 3 columns in word due to data formatting
                        # # issues in further columns
                        # if count <= 2:
                        #     if col_name == 'Year/Country':
                        #         # This IF condition is just to add space to the column name as per Excel sheet to
                        #         # match the names between word and excel. This is an workaround until it is fixed
                        #         col_name = 'Year / Country'
                        if col_name in src_data.columns.values:
                            src_data_final = src_data[col_name]
                            word = []
                            for row in docs.tables[table_count].rows:
                                word.append(row.cells[count].text)

                            src_data_final = [item for item in src_data_final if str(item) != 'nan']
                            # Converting Integer list to String list
                            src_data_final = [str(x) for x in src_data_final]
                            # While reading the table data from word report column name is also being stored. 
                            # Hence removing the column name to get the exact content                            
                            word.pop(0)

                            comparison_result = self.list_comparison_between_reports_data(src_data_final, word)

                            if len(comparison_result) == 0:
                                self.LogScreenshot.fLogScreenshot(message=f"From Sheet '{sheet}', Values in "
                                                                          f"Column '{col_name}' are matching "
                                                                          f"between Expected Excel and Word Report.\n",
                                                                  pass_=True, log=True, screenshot=False)
                            else:
                                self.LogScreenshot.fLogScreenshot(message=f"From Sheet '{sheet}', Values in "
                                                                          f"Column '{col_name}' are not matching "
                                                                          f"between Expected Excel and Word Report.\n "
                                                                          f"Mismatch values are arranged in "
                                                                          f"following order -> Word, Complete Excel "
                                                                          f"and WebExcel Report. {comparison_result}",
                                                                  pass_=False, log=True, screenshot=False)
                                raise Exception("Elements are not matching between Expected Excel "
                                                "and Word Reports")
                            count += 1
                        else:
                            raise Exception("Column names are not matching between Expected Excel "
                                            "and Word Reports")
                    table_count += 1
                else:
                    self.LogScreenshot.fLogScreenshot(message=f"Elements length is not matching between Expected Excel "
                                                              f"and Complete Word Report. Expected Excel Elements "
                                                              f"Length: {len(src_data_len)}\n Word Elements "
                                                              f"Length: {len(word_len)}\n",
                                                      pass_=False, log=True, screenshot=False)
                    raise Exception(f"Elements length is not matching between Expected Excel and "
                                    f"Complete Word Report.")
            except Exception:
                raise Exception("Error in Word report content validation")

    def presencof_publicationtype_col_in_wordreport(self, webexcel_filename, excel_filename, word_filename):
        self.LogScreenshot.fLogScreenshot(message=f"Check presence of Publication Type column in Word Report",
                                          pass_=True, log=True, screenshot=False)
        self.LogScreenshot.fLogScreenshot(message=f"FileNames are: {webexcel_filename}, \n{excel_filename}, "
                                                  f"\n{word_filename}",
                                          pass_=True, log=True, screenshot=False)
        
        # Index of Table number 6 is : 5. Starting point for word table content comparison
        table_counts = {5: "Table 2-2 Clinical study characteristics",
                        6: "Table 2-3 Summary of patient demographics and baseline characteristics",
                        7: "Table 2-4 Efficacy reported in clinical studies",
                        8: "Table 2-5 Safety reported in clinical studies"}
        sheet = 'Clinical Report'

        for table_count, value in table_counts.items():
            webexcel = pd.read_excel(f'ActualOutputs//{webexcel_filename}', sheet_name=sheet, skiprows=3)
            excel = pd.read_excel(f'ActualOutputs//{excel_filename}', sheet_name=sheet, skiprows=3)
            docs = docx.Document(f'ActualOutputs//{word_filename}')
            try:
                table = docs.tables[table_count]
                data = [[cell.text for cell in row.cells] for row in table.rows]
                df_word = pd.DataFrame(data)

                # Using count variable 1 will give the details of 'Publication Type' column in word document
                count = 1
                # 'Publication Type' is the column which will be used for validation
                col_name = 'Publication Type'

                # df_word.values[0] will give the list of column names from the table in Word document
                if col_name in df_word.values[0]:
                    self.LogScreenshot.fLogScreenshot(message=f"Column '{col_name}' is present after 'Short Reference'"
                                                              f" in Word report -> '{value}'",
                                                      pass_=True, log=True, screenshot=False)
                    if col_name in webexcel.columns.values and col_name in excel.columns.values:
                        webex = webexcel[col_name]
                        compex = excel[col_name]
                        word = []
                        for row in docs.tables[table_count].rows:
                            word.append(row.cells[count].text)

                        webex = [item for item in webex if str(item) != 'nan']
                        compex = [item for item in compex if str(item) != 'nan']
                        word.pop(0)

                        if len(webex) == len(compex) == len(word):
                            self.LogScreenshot.fLogScreenshot(message=f"Elements length is matching between Web_Excel, "
                                                                      f"Complete Excel and Complete Word Report. "
                                                                      f"WebExcel Elements Length: {len(webex)}\n "
                                                                      f"Excel Elements Length: {len(compex)}\n "
                                                                      f"Word Elements Length: {len(word)}\n",
                                                              pass_=True, log=True, screenshot=False)      

                            comparison_result = self.list_comparison_between_reports_data(word, compex,
                                                                                          webex_list=webex)

                            if len(comparison_result) == 0:
                                self.LogScreenshot.fLogScreenshot(message=f"From '{value}', Values in Column "
                                                                          f"'{col_name}' are matching between "
                                                                          f"WebExcel, Complete Excel "
                                                                          f"and Word Reports.\n",
                                                                  pass_=True, log=True, screenshot=False)
                            else:
                                self.LogScreenshot.fLogScreenshot(message=f"From '{value}', Values in Column "
                                                                          f"'{col_name}' are not matching between "
                                                                          f"WebExcel, Complete Excel and Word Reports."
                                                                          f"\n Mismatch values are arranged in "
                                                                          f"following order -> Word, Complete Excel "
                                                                          f"and WebExcel Report. {comparison_result}",
                                                                  pass_=False, log=True, screenshot=False)
                                raise Exception("Elements are not matching between Webexcel, Complete Excel and "
                                                "Word Reports")
                        else:
                            self.LogScreenshot.fLogScreenshot(message=f"Elements length is not matching between "
                                                                      f"Web_Excel, Complete Excel and Complete "
                                                                      f"Word Report. WebExcel Elements Length: "
                                                                      f"{len(webex)}\n Excel Elements Length: "
                                                                      f"{len(compex)}\n Word Elements Length: "
                                                                      f"{len(word)}\n",
                                                              pass_=False, log=True, screenshot=False)
                            raise Exception(f"Elements length is not matching between Web_Excel, Complete Excel "
                                            f"and Complete Word Report.")
                    else:
                        raise Exception("Column names are not matching between Webexcel, Complete Excel and "
                                        "Word Reports")
                table_count += 1
            except Exception:
                raise Exception("Error in Word report content validation")

    def check_sorting_order_in_excel_report(self, webexcel_filename, excel_filename):
        self.LogScreenshot.fLogScreenshot(message=f"Check the sorting order in Complete Excel and Standard Excel "
                                                  f"reports", pass_=True, log=True, screenshot=False)
        self.LogScreenshot.fLogScreenshot(message=f"FileNames are: {webexcel_filename} and \n{excel_filename}",
                                          pass_=True, log=True, screenshot=False)
   
        sheet_names = []

        webexcel_data = openpyxl.load_workbook(f'ActualOutputs//{webexcel_filename}')
        excel_data = openpyxl.load_workbook(f'ActualOutputs//{excel_filename}')

        for i in webexcel_data.sheetnames:
            if i in excel_data.sheetnames:
                sheet_names.append(i)

        self.LogScreenshot.fLogScreenshot(message=f"Sheetnames are: {sheet_names}",
                                          pass_=True, log=True, screenshot=False)
        for sheet in sheet_names:
            webexcel = pd.read_excel(f'ActualOutputs//{webexcel_filename}', sheet_name=sheet, skiprows=3)
            compexcel = pd.read_excel(f'ActualOutputs//{excel_filename}', sheet_name=sheet, skiprows=3)

            # Check sorting order from downloaded excel report
            webex_identifier = webexcel["Study Identifier"]
            webex_identifier = [item for item in webex_identifier if str(item) != 'nan']

            # webex_shortreference = webexcel["Short Reference"]
            # webex_shortreference = [item for item in webex_shortreference if str(item) != 'nan']

            compex_identifier = compexcel["Study Identifier"]
            compex_identifier = [item for item in compex_identifier if str(item) != 'nan']

            # compex_shortreference = compexcel["Short Reference"]
            # compex_shortreference = [item for item in compex_shortreference if str(item) != 'nan']

            # Finding duplicate values
            compex_dup_identifier = self.get_duplicates_from_list(compex_identifier)
            # compex_dup_shortref = self.get_duplicates_from_list(compex_shortreference)

            webex_dup_identifier = self.get_duplicates_from_list(webex_identifier)
            # webex_dup_shortref = self.get_duplicates_from_list(webex_shortreference)

            if webex_identifier == sorted(webex_identifier) and compex_identifier == sorted(compex_identifier):
                self.LogScreenshot.fLogScreenshot(message=f"From Sheet '{sheet}', Contents in column 'Study "
                                                          f"Identifier' are in sorted order",
                                                  pass_=True, log=True, screenshot=False)
            else:
                self.LogScreenshot.fLogScreenshot(message=f"From Sheet '{sheet}', Contents in column 'Study "
                                                          f"Identifier' are not in sorted order",
                                                  pass_=False, log=True, screenshot=False)
                raise Exception(f"From Sheet '{sheet}', Contents in column 'Study Identifier' are not in sorted order")
            
            '''Complete Excel Report'''
            # When Study Identifier contains duplicate values then corresponding value in Short Reference
            # column should be in sorted order
            if len(compex_dup_identifier) != 0:
                self.LogScreenshot.fLogScreenshot(
                    message=f"From Complete Excel Report -> '{sheet}' sheet, Duplicate extractions found in "
                            f"'Study Identifier' column, Hence checking the order of corresponding elements in "
                            f"'Short Reference' column. Duplicate Extractions are: '{compex_dup_identifier}'.",
                    pass_=True, log=True, screenshot=False)
                for m in compex_dup_identifier:
                    col_val = compexcel[compexcel["Study Identifier"] == m]
                    col_val_res = col_val["Short Reference"]
                    col_val_res = [item for item in col_val_res if str(item) != 'nan']
                    if col_val_res == sorted(col_val_res):
                        self.LogScreenshot.fLogScreenshot(message=f"For '{m}' in Study Identifier, corresponding "
                                                                  f"contents in column 'Short Reference' are in "
                                                                  f"sorted order in Complete Excel report",
                                                          pass_=True, log=True, screenshot=False)
                    else:
                        self.LogScreenshot.fLogScreenshot(message=f"For '{m}' in Study Identifier, corresponding "
                                                                  f"contents in column 'Short Reference' are not "
                                                                  f"in sorted order in Complete Excel report",
                                                          pass_=True, log=True, screenshot=False)
                        raise Exception(f"For '{m}' in Study Identifier, corresponding contents in column "
                                        f"'Short Reference' are not in sorted order in Complete Excel report")
                
                    compex_dup_shortref = self.get_duplicates_from_list(col_val_res)
                    # When Short Reference contains duplicate values then corresponding value in Publication Type
                    # column should be in sorted order
                    if len(compex_dup_shortref) != 0:
                        self.LogScreenshot.fLogScreenshot(
                            message=f"From Complete Excel Report -> '{sheet}' sheet -> For '{m}' in 'Study "
                                    f"Identifier', Duplicate extractions found in 'Short Reference' column, "
                                    f"Hence checking the order of corresponding elements in 'Publication Type' "
                                    f"column. Duplicate Extractions are: '{compex_dup_shortref}'",
                            pass_=True, log=True, screenshot=False)
                        for n in compex_dup_shortref:
                            first_col_val = compexcel[compexcel["Study Identifier"] == m]
                            col_val = first_col_val[first_col_val["Short Reference"] == n]
                            col_val_res = col_val["Publication Type"]
                            col_val_res = [item for item in col_val_res if str(item) != 'nan']
                            if col_val_res == sorted(col_val_res):
                                self.LogScreenshot.fLogScreenshot(
                                    message=f"For '{m}' in 'Study Identifier' -> '{n}' in Short Reference, "
                                            f"corresponding contents in column 'Publication Type' are in sorted "
                                            f"order in Complete Excel report",
                                    pass_=True, log=True, screenshot=False)
                            else:
                                self.LogScreenshot.fLogScreenshot(
                                    message=f"For '{m}' in 'Study Identifier' -> '{n}' in Short Reference, "
                                            f"corresponding contents in column 'Publication Type' are not in "
                                            f"sorted order in Complete Excel report",
                                    pass_=True, log=True, screenshot=False)
                                raise Exception(f"For '{m}' in 'Study Identifier' -> '{n}' in Short Reference, "
                                                f"corresponding contents in column 'Publication Type' are not in "
                                                f"sorted order in Complete Excel report")

            '''Standard Excel Report'''
            # When Study Identifier contains duplicate values then corresponding value in Short Reference column
            # should be in sorted order
            if len(webex_dup_identifier) != 0:
                self.LogScreenshot.fLogScreenshot(
                    message=f"From Standard Excel Report -> '{sheet}' sheet, Duplicate extractions found in "
                            f"'Study Identifier' column, Hence checking the order of corresponding elements in "
                            f"'Short Reference' column. Duplicate Extractions are: '{webex_dup_identifier}'.",
                    pass_=True, log=True, screenshot=False)
                for m in webex_dup_identifier:
                    col_val = webexcel[webexcel["Study Identifier"] == m]
                    col_val_res = col_val["Short Reference"]
                    col_val_res = [item for item in col_val_res if str(item) != 'nan']
                    if col_val_res == sorted(col_val_res):
                        self.LogScreenshot.fLogScreenshot(message=f"For '{m}' in 'Study Identifier', corresponding "
                                                                  f"contents in column 'Short Reference' are in "
                                                                  f"sorted order in Standard Excel report",
                                                          pass_=True, log=True, screenshot=False)
                    else:
                        self.LogScreenshot.fLogScreenshot(message=f"For '{m}' in 'Study Identifier', corresponding "
                                                                  f"contents in column 'Short Reference' are not in "
                                                                  f"sorted order in Standard Excel report",
                                                          pass_=True, log=True, screenshot=False)
                        raise Exception(f"For '{m}' in 'Study Identifier', corresponding contents in column "
                                        f"'Short Reference' are not in sorted order in Standard Excel report")
                
                    webex_dup_shortref = self.get_duplicates_from_list(col_val_res)
                    # When Short Reference contains duplicate values then corresponding value in Publication Type
                    # column should be in sorted order
                    if len(webex_dup_shortref) != 0:
                        self.LogScreenshot.fLogScreenshot(
                            message=f"From Standard Excel Report -> '{sheet}' sheet -> For '{m}' in 'Study "
                                    f"Identifier', Duplicate extractions found in 'Short Reference' column, "
                                    f"Hence checking the order of corresponding elements in 'Publication Type' "
                                    f"column. Duplicate Extractions are: '{webex_dup_shortref}'",
                            pass_=True, log=True, screenshot=False)
                        for n in webex_dup_shortref:
                            first_col_val = webexcel[webexcel["Study Identifier"] == m]
                            col_val = first_col_val[first_col_val["Short Reference"] == n]
                            col_val_res = col_val["Publication Type"]
                            col_val_res = [item for item in col_val_res if str(item) != 'nan']
                            if col_val_res == sorted(col_val_res):
                                self.LogScreenshot.fLogScreenshot(
                                    message=f"For '{m}' in 'Study Identifier' -> '{n}' in 'Short Reference', "
                                            f"corresponding contents in column 'Publication Type' are in sorted "
                                            f"order in Standard Excel report",
                                    pass_=True, log=True, screenshot=False)
                            else:
                                self.LogScreenshot.fLogScreenshot(
                                    message=f"For '{m}' in 'Study Identifier' -> '{n}' in 'Short Reference', "
                                            f"corresponding contents in column 'Publication Type' are not in "
                                            f"sorted order in Standard Excel report",
                                    pass_=True, log=True, screenshot=False)
                                raise Exception(f"For '{m}' in 'Study Identifier' -> '{n}' in 'Short Reference', "
                                                f"corresponding contents in column 'Publication Type' are not in "
                                                f"sorted order in Standard Excel report")

    def non_oncology_check_sorting_order_in_excel_report(self, webexcel_filename, excel_filename):
        self.LogScreenshot.fLogScreenshot(message=f"Non-Oncology - Check the sorting order in Complete Excel and "
                                                  f"Standard Excel reports", pass_=True, log=True, screenshot=False)
        self.LogScreenshot.fLogScreenshot(message=f"FileNames are: {webexcel_filename} and \n{excel_filename}",
                                          pass_=True, log=True, screenshot=False)
   
        sheet_names = []

        webexcel_data = openpyxl.load_workbook(f'ActualOutputs//{webexcel_filename}')
        excel_data = openpyxl.load_workbook(f'ActualOutputs//{excel_filename}')

        for i in webexcel_data.sheetnames:
            if i in excel_data.sheetnames:
                sheet_names.append(i)

        self.LogScreenshot.fLogScreenshot(message=f"Sheetnames are: {sheet_names}",
                                          pass_=True, log=True, screenshot=False)
        for sheet in sheet_names:
            webexcel = pd.read_excel(f'ActualOutputs//{webexcel_filename}', sheet_name=sheet, skiprows=3)
            compexcel = pd.read_excel(f'ActualOutputs//{excel_filename}', sheet_name=sheet, skiprows=3)

            # Check sorting order from downloaded excel report
            webex_study_id = webexcel["LiveSLR Study ID"]
            webex_study_id = [item for item in webex_study_id if str(item) != 'nan']

            compex_study_id = compexcel["LiveSLR Study ID"]
            compex_study_id = [item for item in compex_study_id if str(item) != 'nan']

            # Finding duplicate values
            compex_dup_study_id = self.get_duplicates_from_list(compex_study_id)

            webex_dup_study_id = self.get_duplicates_from_list(webex_study_id)

            if webex_study_id == sorted(webex_study_id) and compex_study_id == sorted(compex_study_id):
                self.LogScreenshot.fLogScreenshot(
                    message=f"From Sheet '{sheet}', Contents in column 'LiveSLR Study ID' are in sorted order",
                    pass_=True, log=True, screenshot=False)
            else:
                self.LogScreenshot.fLogScreenshot(
                    message=f"From Sheet '{sheet}', Contents in column 'LiveSLR Study ID' are not in sorted order",
                    pass_=False, log=True, screenshot=False)
                raise Exception(f"From Sheet '{sheet}', Contents in column 'LiveSLR Study ID' are not in sorted order")
            
            '''Complete Excel Report'''
            # When LiveSLR Study ID contains duplicate values then corresponding value in Analysis Type
            # column should be in sorted order
            if len(compex_dup_study_id) != 0:
                self.LogScreenshot.fLogScreenshot(
                    message=f"From Complete Excel Report -> '{sheet}' sheet, Duplicate extractions found in "
                            f"'LiveSLR Study ID' column, Hence checking the order of corresponding elements in "
                            f"'Analysis Type' column. Duplicate Extractions are: '{compex_dup_study_id}'.",
                    pass_=True, log=True, screenshot=False)
                for m in compex_dup_study_id:
                    col_val = compexcel[compexcel["LiveSLR Study ID"] == m]
                    col_val_res = col_val["Analysis Type"]
                    col_val_res = [item for item in col_val_res if str(item) != 'nan']

                    col_val_res_final = []
                    # Removing duplicates to get the proper length of duplicate values from Analysis Type column
                    [col_val_res_final.append(x) for x in list(flatten(col_val_res)) if x not in col_val_res_final]

                    if len(col_val_res_final) > 2:
                        if any(["Overall", "Subgroup", "Pooled"] == col_val_res_final[i:i + 3] for i in
                               range(len(col_val_res_final) - 1)):
                            self.LogScreenshot.fLogScreenshot(
                                message=f"For '{m}' in 'LiveSLR Study ID', corresponding contents in column "
                                        f"'Analysis Type' are in sorted order in Complete Excel report",
                                pass_=True, log=True, screenshot=False)
                        else:
                            self.LogScreenshot.fLogScreenshot(
                                message=f"For '{m}' in 'LiveSLR Study ID', corresponding contents in column "
                                        f"'Analysis Type' are not in sorted order in Complete Excel report",
                                pass_=True, log=True, screenshot=False)
                            raise Exception(f"For '{m}' in 'LiveSLR Study ID', corresponding contents in column "
                                            f"'Analysis Type' are not in sorted order in Complete Excel report")
                    elif len(col_val_res_final) > 1:
                        if any(["Overall", "Subgroup"] == col_val_res_final[i:i + 2] for i in
                               range(len(col_val_res_final) - 1)):
                            self.LogScreenshot.fLogScreenshot(
                                message=f"For '{m}' in 'LiveSLR Study ID', corresponding contents in column "
                                        f"'Analysis Type' are in sorted order in Complete Excel report",
                                pass_=True, log=True, screenshot=False)
                        else:
                            self.LogScreenshot.fLogScreenshot(
                                message=f"For '{m}' in 'LiveSLR Study ID', corresponding contents in column "
                                        f"'Analysis Type' are not in sorted order in Complete Excel report",
                                pass_=True, log=True, screenshot=False)
                            raise Exception(f"For '{m}' in 'LiveSLR Study ID', corresponding contents in column "
                                            f"'Analysis Type' are not in sorted order in Complete Excel report")
                    else:
                        if col_val_res == sorted(col_val_res):
                            self.LogScreenshot.fLogScreenshot(
                                message=f"For '{m}' in 'LiveSLR Study ID', corresponding contents in column "
                                        f"'Analysis Type' are in sorted order in Complete Excel report",
                                pass_=True, log=True, screenshot=False)
                        else:
                            self.LogScreenshot.fLogScreenshot(
                                message=f"For '{m}' in 'LiveSLR Study ID', corresponding contents in column "
                                        f"'Analysis Type' are not in sorted order in Complete Excel report",
                                pass_=True, log=True, screenshot=False)
                            raise Exception(f"For '{m}' in 'LiveSLR Study ID', corresponding contents in column "
                                            f"'Analysis Type' are not in sorted order in Complete Excel report")                       
                
                    compex_dup_analysistype = self.get_duplicates_from_list(col_val_res)
                    # When Analysis Type contains duplicate values then corresponding value in Update date (yyyy-mm-dd)
                    # column should be in sorted order
                    if len(compex_dup_analysistype) != 0:
                        self.LogScreenshot.fLogScreenshot(
                            message=f"From Complete Excel Report -> '{sheet}' sheet -> For '{m}' in 'LiveSLR "
                                    f"Study ID', Duplicate extractions found in 'Analysis Type' column, Hence "
                                    f"checking the order of corresponding elements in 'Update date (yyyy-mm-dd)' "
                                    f"column. Duplicate Extractions are: '{compex_dup_analysistype}'",
                            pass_=True, log=True, screenshot=False)
                        for n in compex_dup_analysistype:
                            first_col_val = compexcel[compexcel["LiveSLR Study ID"] == m]
                            col_val = first_col_val[first_col_val["Analysis Type"] == n]
                            col_val_res = col_val["Update date (yyyy-mm-dd)"]
                            col_val_res = [item for item in col_val_res if str(item) != 'nan']
                            if col_val_res == sorted(col_val_res):
                                self.LogScreenshot.fLogScreenshot(
                                    message=f"For '{m}' in 'LiveSLR Study ID' -> '{n}' in 'Analysis Type', "
                                            f"corresponding contents in column 'Update date (yyyy-mm-dd)' are in "
                                            f"sorted order in Complete Excel report",
                                    pass_=True, log=True, screenshot=False)
                            else:
                                self.LogScreenshot.fLogScreenshot(
                                    message=f"For '{m}' in 'LiveSLR Study ID' -> '{n}' in 'Analysis Type', "
                                            f"corresponding contents in column 'Update date (yyyy-mm-dd)' are not "
                                            f"in sorted order in Complete Excel "
                                            f"report", pass_=True, log=True, screenshot=False)
                                raise Exception(f"For '{m}' in 'LiveSLR Study ID' -> '{n}' in 'Analysis Type', "
                                                f"corresponding contents in column 'Update date (yyyy-mm-dd)' are "
                                                f"not in sorted order in Complete Excel report")

            '''Standard Excel Report'''
            # When LiveSLR Study ID contains duplicate values then corresponding value in Analysis Type
            # column should be in sorted order
            if len(webex_dup_study_id) != 0:
                self.LogScreenshot.fLogScreenshot(
                    message=f"From Standard Excel Report -> '{sheet}' sheet, Duplicate extractions found in "
                            f"'LiveSLR Study ID' column, Hence checking the order of corresponding elements in "
                            f"'Analysis Type' column. Duplicate Extractions are: '{webex_dup_study_id}'.",
                    pass_=True, log=True, screenshot=False)
                for m in webex_dup_study_id:
                    col_val = webexcel[webexcel["LiveSLR Study ID"] == m]
                    col_val_res = col_val["Analysis Type"]
                    col_val_res = [item for item in col_val_res if str(item) != 'nan']

                    col_val_res_final = []
                    # Removing duplicates to get the proper length of duplicate values from Analysis Type column
                    [col_val_res_final.append(x) for x in list(flatten(col_val_res)) if x not in col_val_res_final]

                    if len(col_val_res_final) > 2:
                        if any(["Overall", "Subgroup", "Pooled"] == col_val_res_final[i:i + 3] for i in
                               range(len(col_val_res_final) - 1)):
                            self.LogScreenshot.fLogScreenshot(
                                message=f"For '{m}' in 'LiveSLR Study ID', corresponding contents in column "
                                        f"'Analysis Type' are in sorted order in Standard Excel report",
                                pass_=True, log=True, screenshot=False)
                        else:
                            self.LogScreenshot.fLogScreenshot(
                                message=f"For '{m}' in 'LiveSLR Study ID', corresponding contents in column "
                                        f"'Analysis Type' are not in sorted order in Standard Excel report",
                                pass_=True, log=True, screenshot=False)
                            raise Exception(f"For '{m}' in 'LiveSLR Study ID', corresponding contents in column "
                                            f"'Analysis Type' are not in sorted order in Standard Excel report")
                    elif len(col_val_res_final) > 1:
                        if any(["Overall", "Subgroup"] == col_val_res_final[i:i + 2] for i in
                               range(len(col_val_res_final) - 1)):
                            self.LogScreenshot.fLogScreenshot(
                                message=f"For '{m}' in 'LiveSLR Study ID', corresponding contents in column "
                                        f"'Analysis Type' are in sorted order in Standard Excel report",
                                pass_=True, log=True, screenshot=False)
                        else:
                            self.LogScreenshot.fLogScreenshot(
                                message=f"For '{m}' in 'LiveSLR Study ID', corresponding contents in column "
                                        f"'Analysis Type' are not in sorted order in Standard Excel report",
                                pass_=True, log=True, screenshot=False)
                            raise Exception(f"For '{m}' in 'LiveSLR Study ID', corresponding contents in column "
                                            f"'Analysis Type' are not in sorted order in Standard Excel report")
                    else:
                        if col_val_res == sorted(col_val_res):
                            self.LogScreenshot.fLogScreenshot(
                                message=f"For '{m}' in 'LiveSLR Study ID', corresponding contents in column "
                                        f"'Analysis Type' are in sorted order in Standard Excel report",
                                pass_=True, log=True, screenshot=False)
                        else:
                            self.LogScreenshot.fLogScreenshot(
                                message=f"For '{m}' in 'LiveSLR Study ID', corresponding contents in column "
                                        f"'Analysis Type' are not in sorted order in Standard Excel report",
                                pass_=True, log=True, screenshot=False)
                            raise Exception(f"For '{m}' in 'LiveSLR Study ID', corresponding contents in column "
                                            f"'Analysis Type' are not in sorted order in Standard Excel report")                       
                
                    webex_dup_analysistype = self.get_duplicates_from_list(col_val_res)
                    # When Analysis Type contains duplicate values then corresponding value in Update date (yyyy-mm-dd)
                    # column should be in sorted order
                    if len(webex_dup_analysistype) != 0:
                        self.LogScreenshot.fLogScreenshot(
                            message=f"From Standard Excel Report -> '{sheet}' sheet -> For '{m}' in 'LiveSLR "
                                    f"Study ID', Duplicate extractions found in 'Analysis Type' column, Hence "
                                    f"checking the order of corresponding elements in 'Update date (yyyy-mm-dd)' "
                                    f"column. Duplicate Extractions are: '{webex_dup_analysistype}'",
                            pass_=True, log=True, screenshot=False)
                        for n in webex_dup_analysistype:
                            first_col_val = webexcel[webexcel["LiveSLR Study ID"] == m]
                            col_val = first_col_val[first_col_val["Analysis Type"] == n]
                            col_val_res = col_val["Update date (yyyy-mm-dd)"]
                            col_val_res = [item for item in col_val_res if str(item) != 'nan']
                            if col_val_res == sorted(col_val_res):
                                self.LogScreenshot.fLogScreenshot(
                                    message=f"For '{m}' in 'LiveSLR Study ID' -> '{n}' in 'Analysis Type', "
                                            f"corresponding contents in column 'Update date (yyyy-mm-dd)' are in "
                                            f"sorted order in Standard Excel report",
                                    pass_=True, log=True, screenshot=False)
                            else:
                                self.LogScreenshot.fLogScreenshot(
                                    message=f"For '{m}' in 'LiveSLR Study ID' -> '{n}' in 'Analysis Type', "
                                            f"corresponding contents in column 'Update date (yyyy-mm-dd)' are not "
                                            f"in sorted order in Standard Excel report",
                                    pass_=True, log=True, screenshot=False)
                                raise Exception(f"For '{m}' in 'LiveSLR Study ID' -> '{n}' in 'Analysis Type', "
                                                f"corresponding contents in column 'Update date (yyyy-mm-dd)' are "
                                                f"not in sorted order in Standard Excel report")

    def test_prisma_ele_comparison_between_Excel_and_Word_Report(self, pop_data, slr_type, add_criteria, filepath, env):

        # Go to live slr page
        # self.go_to_page("SLR_Homepage", env)
        # self.click("liveslr_reset_filter", env)
        self.refreshpage()
        time.sleep(2)
        self.select_data(f"{pop_data[0][0]}", f"{pop_data[0][1]}", env)
        self.select_data(f"{slr_type[0][0]}", f"{slr_type[0][1]}", env)
        self.select_sub_section(f"{add_criteria[0][0]}", f"{add_criteria[0][1]}", env, f"{add_criteria[0][2]}")
        self.select_sub_section(f"{add_criteria[1][0]}", f"{add_criteria[1][1]}", env, f"{add_criteria[1][2]}")
        self.select_sub_section(f"{add_criteria[2][0]}", f"{add_criteria[2][1]}", env, f"{add_criteria[2][2]}")
        self.select_sub_section(f"{add_criteria[3][0]}", f"{add_criteria[3][1]}", env, f"{add_criteria[3][2]}")

        self.generate_download_report("excel_report", env)
        # time.sleep(5)
        # excel_filename = self.getFilenameAndValidate(180)
        # excel_filename = self.get_latest_filename(UnivWaitFor=180)
        excel_filename = self.get_and_validate_filename(filepath)

        self.generate_download_report("word_report", env)
        # time.sleep(5)
        # word_filename = self.getFilenameAndValidate(180)
        # word_filename = self.get_latest_filename(UnivWaitFor=180)
        word_filename = self.get_and_validate_filename(filepath)

        excel = pd.read_excel(f'ActualOutputs//{excel_filename}', sheet_name='Updated PRISMA', usecols='C', skiprows=11)
        excel_studies_col = excel['Unique studies']
        # Removing NAN values
        excel_studies_col = [item for item in excel_studies_col if str(item) != 'nan']
        # Converting list values from float to int
        excel_studies_col = [int(x) for x in excel_studies_col]

        # Reading PRISMA Value from Complete Word Report
        docs = docx.Document(f'ActualOutputs//{word_filename}')
        word = []
        for row in docs.tables[4].rows:
            word.append(row.cells[1].text)
        # Removing unwanted values
        word.pop(2)
        word.pop(0)
        # Converting list values from str to int
        word = [int(x) for x in word]

        if excel_studies_col == word:
            self.LogScreenshot.fLogScreenshot(message=f"'Updated Prisma' count is matching between Excel and Word "
                                                      f"report. Excel report 'Updated Prisma' tab count is "
                                                      f"{excel_studies_col} and Word Report 'Updated Prisma' "
                                                      f"table count is {word}",
                                              pass_=True, log=True, screenshot=False)
        else:
            self.LogScreenshot.fLogScreenshot(message=f"'Updated Prisma' count is not matching between Excel and "
                                                      f"Word report. Excel report 'Updated Prisma' tab count is "
                                                      f"{excel_studies_col} and Word Report 'Updated Prisma' "
                                                      f"table count is {word}",
                                              pass_=False, log=True, screenshot=False)
            raise Exception(f"'Updated Prisma' count is not matching between Excel and Word report. Excel report "
                            f"'Updated Prisma' tab count is {excel_studies_col} and Word Report 'Updated Prisma' "
                            f"table count is {word}")

    def prisma_ele_comparison_between_Excel_and_UI(self, locatorname, pop_data, slr_type, add_criteria, sheet,
                                                        filepath, env, prj_name):

        # Read expected categoris from data sheet
        expected_prisma_count = self.exbase.get_individual_col_data(filepath, locatorname, 'Sheet1',
                                                                    'ExpectedPrismaCount')

        # Converting list values from float to int
        expected_prisma_count = [int(x) for x in expected_prisma_count]      

        # Go to live slr page
        self.go_to_page("SLR_Homepage", env)
        # self.click("liveslr_reset_filter", env)
        self.refreshpage()
        time.sleep(2)
        self.select_data(f"{pop_data[0][0]}", f"{pop_data[0][1]}", env)
        self.select_data(f"{slr_type[0][0]}", f"{slr_type[0][1]}", env)
        for k in add_criteria:
            self.select_sub_section(f"{k[0]}", f"{k[1]}", env, f"{k[2]}")        

        self.scroll("New_total_selected", env)
        locators = ['Original_SLR_Count', 'Sub_Pop_Count', 'Line_of_Therapy_Count', 'Intervention_Count',
                    'Study_Design_Count', 'Reported_Var_Count', 'New_total_selected']
        ui_add_critera_values = []
        for i in locators:
            ui_add_critera_values.append(self.get_text(i, env))
        # Converting list values from str to int
        ui_add_critera_values = [int(x) for x in ui_add_critera_values]

        self.LogScreenshot.fLogScreenshot(message=f"Expected PRISMA Count Values are: "
                                                  f"Original SLR Count : {expected_prisma_count[0]}, Sub Pop Count : "
                                                  f"{expected_prisma_count[1]}, LOT Count : {expected_prisma_count[2]},"
                                                  f" Intervention Count : {expected_prisma_count[3]}, "
                                                  f"Study Design Count : {expected_prisma_count[4]}, Reported "
                                                  f"Variable Count : {expected_prisma_count[5]}, Total Selected "
                                                  f"Count : {expected_prisma_count[6]}",
                                          pass_=True, log=True, screenshot=False)

        self.LogScreenshot.fLogScreenshot(message=f"Actual PRISMA Count Values are: "
                                                  f"Original SLR Count : {ui_add_critera_values[0]}, Sub Pop Count : "
                                                  f"{ui_add_critera_values[1]}, LOT Count : {ui_add_critera_values[2]},"
                                                  f" Intervention Count : {ui_add_critera_values[3]}, "
                                                  f"Study Design Count : {ui_add_critera_values[4]}, Reported "
                                                  f"Variable Count : {ui_add_critera_values[5]}, Total Selected "
                                                  f"Count : {ui_add_critera_values[6]}",
                                          pass_=True, log=True, screenshot=True)

        self.generate_download_report("excel_report", env)
        excel_filename = self.get_and_validate_filename(filepath)

        '''This condition is writted due to LIVEHTA-2473 implementation for Non-Oncology projects'''
        if prj_name == "Oncology":
            excel = pd.read_excel(f'ActualOutputs//{excel_filename}', sheet_name=sheet, usecols='C', skiprows=11)
        else:
            excel = pd.read_excel(f'ActualOutputs//{excel_filename}', sheet_name=sheet, usecols='C', skiprows=12)
        excel_studies_col = excel['Unique studies']
        # Removing NAN values
        excel_studies_col = [item for item in excel_studies_col if str(item) != 'nan']
        # Converting list values from float to int
        excel_studies_col = [int(x) for x in excel_studies_col]

        prisma_count_comparison = self.list_comparison_between_reports_data(expected_prisma_count, excel_studies_col,
                                                                            ui_add_critera_values)

        if len(prisma_count_comparison) == 0:
            self.LogScreenshot.fLogScreenshot(message=f"Count seen in excel report under 'Updated Prisma tab' and "
                                                      f"'Updated PRISMA table' Count in UI is matching with Expected "
                                                      f"PRISMA Count.", pass_=True, log=True, screenshot=False)
        else:
            self.LogScreenshot.fLogScreenshot(message=f"Count seen in excel report under 'Updated Prisma tab' and "
                                                      f"'Updated PRISMA table' Count in UI is not matching with "
                                                      f"Expected PRISMA Count. Mismatch values are arranged in "
                                                      f"following order -> Expected count, Excel count, UI count. "
                                                      f"{prisma_count_comparison}",
                                              pass_=False, log=True, screenshot=False)
            raise Exception(f"Count seen in excel report under 'Updated Prisma tab' and 'Updated PRISMA table' "
                            f"Count in UI is not matching with Expected PRISMA Count")

    def test_prisma_count_comparison_between_prismatab_and_excludedstudiesliveslr(self, pop_data, slr_type,
                                                                                  add_criteria, filepath, env):

        # Go to live slr page
        # self.go_to_page("SLR_Homepage", env)
        # self.click("liveslr_reset_filter", env)
        self.refreshpage()
        time.sleep(2)
        self.select_data(f"{pop_data[0][0]}", f"{pop_data[0][1]}", env)
        self.select_data(f"{slr_type[0][0]}", f"{slr_type[0][1]}", env)
        self.select_sub_section(f"{add_criteria[0][0]}", f"{add_criteria[0][1]}", env, f"{add_criteria[0][2]}")
        self.select_sub_section(f"{add_criteria[1][0]}", f"{add_criteria[1][1]}", env, f"{add_criteria[1][2]}")
        self.select_sub_section(f"{add_criteria[2][0]}", f"{add_criteria[2][1]}", env, f"{add_criteria[2][2]}")
        self.select_sub_section(f"{add_criteria[3][0]}", f"{add_criteria[3][1]}", env, f"{add_criteria[3][2]}")

        self.generate_download_report("excel_report", env)
        # time.sleep(5)
        # excel_filename = self.getFilenameAndValidate(180)
        # excel_filename = self.get_latest_filename(UnivWaitFor=180)
        excel_filename = self.get_and_validate_filename(filepath)

        prisma_tab = pd.read_excel(f'ActualOutputs//{excel_filename}', sheet_name='Updated PRISMA', usecols='C',
                                   skiprows=11)
        excludedstudies_liveslr_tab = pd.read_excel(f'ActualOutputs//{excel_filename}',
                                                    sheet_name='Excluded studies - LiveSLR', skiprows=2)
        prisma_tab_col = prisma_tab['Unique studies']
        # Removing NAN values
        prisma_tab_col = [item for item in prisma_tab_col if str(item) != 'nan']
        # Converting list values from float to int
        prisma_tab_col = [int(x) for x in prisma_tab_col]

        # Get Unique Study Identifier value from corresponding to Reason for Rejection values
        excludedstudies_liveslr = excludedstudies_liveslr_tab["Reason for Rejection"]
        excludedstudies_liveslr = [item for item in excludedstudies_liveslr if str(item) != 'nan']
        excludedstudies_liveslr_final = []
        # Removing the duplicates
        [excludedstudies_liveslr_final.append(x) for x in list(flatten(excludedstudies_liveslr))
         if x not in excludedstudies_liveslr_final]

        self.LogScreenshot.fLogScreenshot(message=f"Unique Reason for Rejection column values are : "
                                                  f"{excludedstudies_liveslr_final}",
                                          pass_=True, log=True, screenshot=False)

        for m in excludedstudies_liveslr_final:
            unique_study_identifier = []
            col_val = excludedstudies_liveslr_tab[excludedstudies_liveslr_tab["Reason for Rejection"] == m]
            study_identifier = col_val["Study Identifier"]
            study_identifier = [item for item in study_identifier if str(item) != 'nan']
            # Removing duplicates to get the proper length of Study Identifier data
            [unique_study_identifier.append(x) for x in list(flatten(study_identifier))
             if x not in unique_study_identifier]
            if len(unique_study_identifier) in prisma_tab_col[1:6]:
                self.LogScreenshot.fLogScreenshot(message=f"From 'Excluded studies - LiveSLR' sheet -> for '{m}' "
                                                          f"Reason for Rejection the Selected studies(Unique Study "
                                                          f"Identifier)count is matching with the selected count in "
                                                          f"Updated PRISMA tab",
                                                  pass_=True, log=True, screenshot=False)
            else:
                self.LogScreenshot.fLogScreenshot(message=f"From 'Excluded studies - LiveSLR' sheet -> for '{m}' "
                                                          f"Reason for Rejection the Selected studies(Unique Study "
                                                          f"Identifier)count is not matching with the selected count "
                                                          f"in Updated PRISMA tab",
                                                  pass_=False, log=True, screenshot=False)
                raise Exception(f"From 'Excluded studies - LiveSLR' sheet -> for '{m}' Reason for Rejection the "
                                f"Selected studies(Unique Study Identifier)count is not matching with the selected "
                                f"count in Updated PRISMA tab")

    def test_prisma_tab_format_changes(self, pop_data, slr_type, add_criteria, filepath, env):

        # Go to live slr page
        # self.go_to_page("SLR_Homepage", env)
        # self.click("liveslr_reset_filter", env)
        self.refreshpage()
        time.sleep(2)
        self.select_data(f"{pop_data[0][0]}", f"{pop_data[0][1]}", env)
        self.select_data(f"{slr_type[0][0]}", f"{slr_type[0][1]}", env)
        self.select_sub_section(f"{add_criteria[0][0]}", f"{add_criteria[0][1]}", env, f"{add_criteria[0][2]}")
        self.select_sub_section(f"{add_criteria[1][0]}", f"{add_criteria[1][1]}", env, f"{add_criteria[1][2]}")
        self.select_sub_section(f"{add_criteria[2][0]}", f"{add_criteria[2][1]}", env, f"{add_criteria[2][2]}")
        self.select_sub_section(f"{add_criteria[3][0]}", f"{add_criteria[3][1]}", env, f"{add_criteria[3][2]}")

        self.generate_download_report("excel_report", env)
        # time.sleep(5)
        # excel_filename = self.getFilenameAndValidate(180)
        # excel_filename = self.get_latest_filename(UnivWaitFor=180)
        excel_filename = self.get_and_validate_filename(filepath)

        # Read PRISMA value from Complete Excel Sheet
        excel = openpyxl.load_workbook(f'ActualOutputs//{excel_filename}')
        excel_sheet = excel['Updated PRISMA']
        row_13 = excel_sheet['B13'].value
        row_12 = excel_sheet['C12'].value
        row_14 = excel_sheet['B14'].value
        row_15 = excel_sheet['B15'].value
        row_21 = excel_sheet['B21'].value
        row_23 = excel_sheet['B23'].value
        
        if "From original SLR" in row_13:
            self.LogScreenshot.fLogScreenshot(message=f"Format has been updated from 'Original SLR' to 'From "
                                                      f"original SLR'",
                                              pass_=True, log=True, screenshot=False)
        else:
            self.LogScreenshot.fLogScreenshot(message=f"Format has not been updated from 'Original SLR' to 'From "
                                                      f"original SLR'",
                                              pass_=False, log=True, screenshot=False)
            raise Exception(f"Format has not been updated from 'Original SLR' to 'From original SLR'")

        if "Unique studies" in row_12:
            self.LogScreenshot.fLogScreenshot(message=f"Column Title has been added with name 'Unique studies'",
                                              pass_=True, log=True, screenshot=False)
        else:
            self.LogScreenshot.fLogScreenshot(message=f"Column Title has not been added with name 'Unique studies'",
                                              pass_=False, log=True, screenshot=False)
            raise Exception(f"Column Title has not been added with name 'Unique studies'")

        if row_14 is None:
            self.LogScreenshot.fLogScreenshot(message=f"An empty row below 'From original SLR' has been added",
                                              pass_=True, log=True, screenshot=False)
        else:
            self.LogScreenshot.fLogScreenshot(message=f"An empty row below 'From original SLR' has not been added",
                                              pass_=False, log=True, screenshot=False)
            raise Exception(f"An empty row below 'From original SLR' has not been added")

        if "Excluded based on" in row_15:
            self.LogScreenshot.fLogScreenshot(message=f"An additional row has been added above 'sub-population' and "
                                                      f"the column title is 'Excluded based on'",
                                              pass_=True, log=True, screenshot=False)
        else:
            self.LogScreenshot.fLogScreenshot(message=f"An additional row has not been added above 'sub-population'",
                                              pass_=False, log=True, screenshot=False)
            raise Exception(f"An additional row has not been added above 'sub-population'")

        if row_21 is None:
            self.LogScreenshot.fLogScreenshot(message=f"An empty row above 'New total selected' has been added",
                                              pass_=True, log=True, screenshot=False)
        else:
            self.LogScreenshot.fLogScreenshot(message=f"An empty row above 'New total selected' has not been added",
                                              pass_=False, log=True, screenshot=False)
            raise Exception(f"An empty row above 'New total selected' has not been added")

        if "Note: LiveSLR counts the number of original studies, however, for each original study, there might be " \
           "multiple extractions covering original or sub-group patient populations. Therefore, the number of the " \
           "new total selected and the number of excluded studies may not add up to be the total original SLR count." \
           " Please contact your Cytel SLR project lead if you require more assistance on this." in row_23:
            self.LogScreenshot.fLogScreenshot(message=f"Static text '{row_23}' has been added.",
                                              pass_=True, log=True, screenshot=False)
        else:
            self.LogScreenshot.fLogScreenshot(message=f"Static text '{row_23}' has not been added.",
                                              pass_=False, log=True, screenshot=False)
            raise Exception(f"Static text '{row_23}' has not been added.")                                            

    def test_interventional_to_clinical_changes(self, filepath, pop_list, env):

        # Go to live slr page
        self.go_to_page("SLR_Homepage", env)
        self.select_data(pop_list[0][0], pop_list[0][1], env)
        time.sleep(1)
        type_of_slr_text = self.get_text("slrtype_first_option_text", env)
        self.select_data("Clinical", "Clinical_radio_button", env)

        self.click("NMA_Button", env)
        nma_param_clinical_list = self.get_text("NMA_parameters_list_clinical", env)

        self.generate_download_report("excel_report", env)
        excel_filename = self.get_and_validate_filename(filepath)

        self.generate_download_report("word_report", env)
        word_filename = self.get_and_validate_filename(filepath)

        self.preview_result("preview_results", env)
        self.table_display_check("Table", env)
        web_table_title = self.get_text("web_table_title", env)
        self.generate_download_report("Export_as_excel", env)
        webexcel_filename = self.get_and_validate_filename(filepath)
        self.back_to_report_page("Back_to_search_page", env)

        if search("Clinical", type_of_slr_text) and not search("Interventional", type_of_slr_text):
            self.LogScreenshot.fLogScreenshot(message=f"From 'Search LiveSLR' page -> Under Select Type of SLR first "
                                                      f"option is updated from Interventional to 'Clinical'",
                                              pass_=True, log=True, screenshot=False)
        else:
            self.LogScreenshot.fLogScreenshot(message=f"From 'Search LiveSLR' page -> Under Select Type of SLR first "
                                                      f"option is not updated from Interventional to 'Clinical'. "
                                                      f"Actual value is {type_of_slr_text}",
                                              pass_=False, log=True, screenshot=False)
            raise Exception(f"From 'Search LiveSLR' page -> Under Select Type of SLR first option is not updated "
                            f"from Interventional to 'Clinical'")
                
        if search("Clinical", nma_param_clinical_list) and not search("Interventional", nma_param_clinical_list):
            self.LogScreenshot.fLogScreenshot(message=f"From 'Search LiveSLR' page -> Go to 'Select Data for  NMA' -> "
                                                      f"Under select NMA parameters section is updated from "
                                                      f"Interventional to 'Clinical'",
                                              pass_=True, log=True, screenshot=False)
        else:
            self.LogScreenshot.fLogScreenshot(message=f"From 'Search LiveSLR' page -> Go to 'Select Data for  NMA' -> "
                                                      f"Under select NMA parameters section is not updated from "
                                                      f"Interventional to 'Clinical'",
                                              pass_=False, log=True, screenshot=False)
            raise Exception(f"From 'Search LiveSLR' page -> Go to 'Select Data for  NMA' -> Under select NMA "
                            f"parameters section is not updated from Interventional to 'Clinical'")

        self.click("protocol_link", env)
        pages = [['picos', 'picos_pop_dropdown', 'picos_study_type_dropdown'],
                 ['searchstrategy', 'searchstrategy_pop_dropdown', 'searchstrategy_study_type_dropdown'],
                 ['prismas', 'prisma_pop_dropdown', 'prisma_study_type_dropdown']]
        for i in pages:
            self.click(i[0], env)
            pop_ele = self.select_element(i[1], env)
            select1 = Select(pop_ele)
            select1.select_by_visible_text("NewImportLogic_1 - Test_Automation_1")

            stdy_ele = self.select_element(i[2], env)
            select2 = Select(stdy_ele)
            opts = select2.options
            dropdown_li = []
            for j in opts:
                dropdown_li.append(j.text)
            
            if "Clinical" in dropdown_li and "Interventional" not in dropdown_li:
                self.LogScreenshot.fLogScreenshot(message=f"From page '{i[0]}' -> Under Study type dropdown "
                                                          f"Interventional is changed to 'Clinical'",
                                                  pass_=True, log=True, screenshot=False)
            else:
                self.LogScreenshot.fLogScreenshot(message=f"From page '{i[0]}' -> Under Study type dropdown "
                                                          f"Interventional is not changed to 'Clinical'",
                                                  pass_=False, log=True, screenshot=False)
                raise Exception(f"From page '{i[0]}' -> Under Study type dropdown Interventional is not changed "
                                f"to 'Clinical'")

        self.click("manage_qa_data_button", env)
        time.sleep(1)
        pop_ele = self.select_element("select_pop_dropdown", env)
        select1 = Select(pop_ele)
        select1.select_by_visible_text("NewImportLogic_1 - Test_Automation_1")

        stdy_ele = self.select_element("select_stdy_type_dropdown", env)
        select2 = Select(stdy_ele)
        opts = select2.options
        qa_data_dropdown_li = []
        for m in opts:
            qa_data_dropdown_li.append(m.text)
        
        if "Clinical" in qa_data_dropdown_li and "Interventional" not in qa_data_dropdown_li:
            self.LogScreenshot.fLogScreenshot(message=f"From 'Manage QA Data' page -> Under Study type dropdown "
                                                      f"Interventional is changed to 'Clinical'",
                                              pass_=True, log=True, screenshot=False)
        else:
            self.LogScreenshot.fLogScreenshot(message=f"From 'Manage QA Data' page -> Under Study type dropdown "
                                                      f"Interventional is not changed to 'Clinical'",
                                              pass_=False, log=True, screenshot=False)
            raise Exception(f"From 'Manage QA Data' page -> Under Study type dropdown Interventional is not changed "
                            f"to 'Clinical'")
        
        if search("Clinical", web_table_title) and not search("Interventional", web_table_title):
            self.LogScreenshot.fLogScreenshot(message=f"WebExcel Report -> Title of 'Preview Results' section is "
                                                      f"updated from Interventional to 'Clinical'",
                                              pass_=True, log=True, screenshot=False)
        else:
            self.LogScreenshot.fLogScreenshot(message=f"WebExcel Report -> Title of 'Preview Results' section is not "
                                                      f"updated from Interventional to 'Clinical'",
                                              pass_=False, log=True, screenshot=False)
            raise Exception(f"WebExcel Report -> Title of 'Preview Results' section is not updated from "
                            f"Interventional to 'Clinical'")
        
        # Read value from Complete Excel and WebExcel
        excel = openpyxl.load_workbook(f'ActualOutputs//{excel_filename}')
        webexcel = openpyxl.load_workbook(f'ActualOutputs//{webexcel_filename}')
        excel_sheetnames = excel.sheetnames
        webexcel_sheetnames = webexcel.sheetnames

        webex_sheet = webexcel.active
        webex_title = webex_sheet['A1'].value
        if search("Clinical", webex_title) and not search("Interventional", webex_title) and "Clinical Report" in \
                webexcel_sheetnames:
            self.LogScreenshot.fLogScreenshot(message=f"WebExcel Report -> Title and sheet name is updated from "
                                                      f"Interventional to 'Clinical'",
                                              pass_=True, log=True, screenshot=False)
        else:
            self.LogScreenshot.fLogScreenshot(message=f"WebExcel Report -> Title and sheet name is not updated from "
                                                      f"Interventional to 'Clinical'",
                                              pass_=False, log=True, screenshot=False)
            raise Exception(f"WebExcel Report -> Title and sheet name is not updated from Interventional to 'Clinical'")

        excel_sheet_toc = excel['TOC']
        toc_heading = excel_sheet_toc['A1'].value
        toc_content_6 = excel_sheet_toc['B10'].value

        if search("Clinical", toc_heading) and not search("Interventional", toc_heading):
            self.LogScreenshot.fLogScreenshot(message=f"TOC sheet heading is updated from Interventional to 'Clinical'",
                                              pass_=True, log=True, screenshot=False)
        else:
            self.LogScreenshot.fLogScreenshot(message=f"TOC sheet heading is not updated from Interventional to "
                                                      f"'Clinical'",
                                              pass_=False, log=True, screenshot=False)
            raise Exception(f"TOC sheet heading is not updated from Interventional to 'Clinical'")

        if search("Clinical", toc_content_6) and not search("Interventional", toc_content_6):
            self.LogScreenshot.fLogScreenshot(message=f"TOC sheet -> table of content number 6 is updated from "
                                                      f"Interventional to 'Clinical'",
                                              pass_=True, log=True, screenshot=False)
        else:
            self.LogScreenshot.fLogScreenshot(message=f"TOC sheet -> table of content number 6 is not updated from "
                                                      f"Interventional to 'Clinical'",
                                              pass_=False, log=True, screenshot=False)
            raise Exception(f"TOC sheet -> table of content number 6 is not updated from Interventional to 'Clinical'")

        excel_sheet_picos = excel['PICOS']
        picos_heading = excel_sheet_picos['A3'].value

        if search("CLINICAL", picos_heading) and not search("INTERVENTIONAL", picos_heading):
            self.LogScreenshot.fLogScreenshot(message=f"PICOS sheet heading is updated from INTERVENTIONAL to "
                                                      f"'CLINICAL'",
                                              pass_=True, log=True, screenshot=False)
        else:
            self.LogScreenshot.fLogScreenshot(message=f"PICOS sheet heading is not updated from INTERVENTIONAL to "
                                                      f"'CLINICAL'",
                                              pass_=False, log=True, screenshot=False)
            raise Exception(f"PICOS sheet heading is not updated from INTERVENTIONAL to 'CLINICAL'")

        excel_sheet_inc_exc = excel['INC-EXC']
        inc_sheet_heading = excel_sheet_inc_exc['A1'].value
        inc_table_heading = excel_sheet_inc_exc['A4'].value

        if search("Clinical", inc_sheet_heading) and not search("Interventional", inc_sheet_heading):
            self.LogScreenshot.fLogScreenshot(message=f"INCLUSION AND EXCLUSION CRITERIA sheet heading is updated "
                                                      f"from Interventional to 'Clinical'",
                                              pass_=True, log=True, screenshot=False)
        else:
            self.LogScreenshot.fLogScreenshot(message=f"INCLUSION AND EXCLUSION CRITERIA sheet heading is not "
                                                      f"updated from Interventional to 'Clinical'",
                                              pass_=False, log=True, screenshot=False)
            raise Exception(f"INCLUSION AND EXCLUSION CRITERIA sheet heading is not updated from Interventional to "
                            f"'Clinical'")

        if search("Clinical", inc_table_heading) and not search("Interventional", inc_table_heading):
            self.LogScreenshot.fLogScreenshot(message=f"INCLUSION AND EXCLUSION CRITERIA table heading is updated "
                                                      f"from Interventional to 'Clinical'",
                                              pass_=True, log=True, screenshot=False)
        else:
            self.LogScreenshot.fLogScreenshot(message=f"INCLUSION AND EXCLUSION CRITERIA table heading is not updated "
                                                      f"from Interventional to 'Clinical'",
                                              pass_=False, log=True, screenshot=False)
            raise Exception(f"INCLUSION AND EXCLUSION CRITERIA table heading is not updated from Interventional "
                            f"to 'Clinical'")

        excel_sheet_strategy = excel['SEARCH STRATEGIES']
        strategy_sheet_heading = excel_sheet_strategy['A1'].value
        strategy_table_heading = excel_sheet_strategy['A3'].value

        if search("Clinical", strategy_sheet_heading) and not search("Interventional", strategy_sheet_heading):
            self.LogScreenshot.fLogScreenshot(message=f"SEARCH STRATEGY sheet heading is updated from Interventional "
                                                      f"to 'Clinical'",
                                              pass_=True, log=True, screenshot=False)
        else:
            self.LogScreenshot.fLogScreenshot(message=f"SEARCH STRATEGY sheet heading is not updated from "
                                                      f"Interventional to 'Clinical'",
                                              pass_=False, log=True, screenshot=False)
            raise Exception(f"SEARCH STRATEGY sheet heading is not updated from Interventional to 'Clinical'")

        if search("Clinical", strategy_table_heading) and not search("Interventional", strategy_table_heading):
            self.LogScreenshot.fLogScreenshot(message=f"SEARCH STRATEGY table heading is updated from Interventional "
                                                      f"to 'Clinical'",
                                              pass_=True, log=True, screenshot=False)
        else:
            self.LogScreenshot.fLogScreenshot(message=f"SEARCH STRATEGY table heading is not updated from "
                                                      f"Interventional to 'Clinical'",
                                              pass_=False, log=True, screenshot=False)
            raise Exception(f"SEARCH STRATEGY table heading is not updated from Interventional to 'Clinical'")

        excel_sheet_report = excel['Clinical Report']
        report_sheet_heading = excel_sheet_report['A1'].value 

        if "Clinical Report" in excel_sheetnames and "Interventional Report" not in excel_sheetnames:
            self.LogScreenshot.fLogScreenshot(message=f"Report sheet name is updated from Interventional to 'Clinical'",
                                              pass_=True, log=True, screenshot=False)
        else:
            self.LogScreenshot.fLogScreenshot(message=f"Report sheet name is not updated from Interventional to "
                                                      f"'Clinical'",
                                              pass_=False, log=True, screenshot=False)
            raise Exception(f"Report sheet name is not updated from Interventional to 'Clinical'")

        if search("Clinical", report_sheet_heading) and not search("Interventional", report_sheet_heading):
            self.LogScreenshot.fLogScreenshot(message=f"Report sheet heading is updated from Interventional to "
                                                      f"'Clinical'",
                                              pass_=True, log=True, screenshot=False)
        else:
            self.LogScreenshot.fLogScreenshot(message=f"Report sheet heading is not updated from Interventional "
                                                      f"to 'Clinical'",
                                              pass_=False, log=True, screenshot=False)
            raise Exception(f"Report sheet heading is not updated from Interventional to 'Clinical'")

    def test_publication_identifier_count_in_updated_prisma_tab(self, pop_data, slr_type, add_criteria, filepath, env):

        # Go to live slr page
        # self.go_to_page("SLR_Homepage", env)
        # self.click("liveslr_reset_filter", env)
        self.refreshpage()
        self.select_data(f"{pop_data[0][0]}", f"{pop_data[0][1]}", env)
        self.select_data(f"{slr_type[0][0]}", f"{slr_type[0][1]}", env)
        self.select_sub_section(f"{add_criteria[0][0]}", f"{add_criteria[0][1]}", env, f"{add_criteria[0][2]}")
        self.select_sub_section(f"{add_criteria[1][0]}", f"{add_criteria[1][1]}", env, f"{add_criteria[1][2]}")
        self.select_sub_section(f"{add_criteria[2][0]}", f"{add_criteria[2][1]}", env, f"{add_criteria[2][2]}")
        self.select_sub_section(f"{add_criteria[3][0]}", f"{add_criteria[3][1]}", env, f"{add_criteria[3][2]}")

        self.generate_download_report("excel_report", env)
        # time.sleep(5)
        # excel_filename = self.getFilenameAndValidate(180)
        # excel_filename = self.get_latest_filename(UnivWaitFor=180)
        excel_filename = self.get_and_validate_filename(filepath)

        prisma_tab = pd.read_excel(f'ActualOutputs//{excel_filename}', sheet_name='Updated PRISMA', usecols='D',
                                   skiprows=11)
        excludedstudies_liveslr_tab = pd.read_excel(f'ActualOutputs//{excel_filename}',
                                                    sheet_name='Excluded studies - LiveSLR', skiprows=2)
        prisma_tab_col = prisma_tab['Publications']
        # Removing NAN values
        prisma_tab_col = [item for item in prisma_tab_col if str(item) != 'nan']
        # Converting list values from float to int
        prisma_tab_col = [int(x) for x in prisma_tab_col]

        # Get Unique Study Identifier value from corresponding to Reason for Rejection values
        excludedstudies_liveslr = excludedstudies_liveslr_tab["Reason for Rejection"]
        excludedstudies_liveslr = [item for item in excludedstudies_liveslr if str(item) != 'nan']
        excludedstudies_liveslr_final = []
        # Removing the duplicates
        [excludedstudies_liveslr_final.append(x) for x in list(flatten(excludedstudies_liveslr))
         if x not in excludedstudies_liveslr_final]

        self.LogScreenshot.fLogScreenshot(message=f"Unique Reason for Rejection column values are : "
                                                  f"{excludedstudies_liveslr_final}",
                                          pass_=True, log=True, screenshot=False)

        for m in excludedstudies_liveslr_final:
            unique_pub_identifier = []
            col_val = excludedstudies_liveslr_tab[excludedstudies_liveslr_tab["Reason for Rejection"] == m]
            pub_identifier = col_val["Publication Identifier"]
            pub_identifier = [item for item in pub_identifier if str(item) != 'nan']
            # Checking if list is nested or not
            if any(isinstance(i, str) for i in pub_identifier):
                # Converting a list containing String type values into list/tuple type
                res = []
                for xy in pub_identifier:
                    res.append(ast.literal_eval(xy))
                # Removing duplicates to get the proper length of Study Identifier data
                [unique_pub_identifier.append(x) for x in list(flatten(res))
                 if x not in unique_pub_identifier]
            else:
                # Removing duplicates to get the proper length of Study Identifier data
                [unique_pub_identifier.append(x) for x in list(flatten(pub_identifier))
                 if x not in unique_pub_identifier]

            if len(unique_pub_identifier) in prisma_tab_col[1:6]:
                self.LogScreenshot.fLogScreenshot(message=f"From 'Excluded studies - LiveSLR' sheet -> for '{m}' "
                                                          f"Reason for Rejection the unique Publication "
                                                          f"Identifier count is matching with the count in "
                                                          f"Updated PRISMA tab -> 'Publications' column. The count "
                                                          f"value is : {len(unique_pub_identifier)}",
                                                  pass_=True, log=True, screenshot=False)
            else:
                self.LogScreenshot.fLogScreenshot(message=f"From 'Excluded studies - LiveSLR' sheet -> for '{m}' "
                                                          f"Reason for Rejection the unique Publication "
                                                          f"Identifier count is not matching with the count in "
                                                          f"Updated PRISMA tab -> 'Publications' column. The count "
                                                          f"value is : {len(unique_pub_identifier)}",
                                                  pass_=False, log=True, screenshot=False)
                raise Exception(f"From 'Excluded studies - LiveSLR' sheet -> for '{m}' Reason for Rejection the "
                                f"unique Publication Identifier count is not matching with the "
                                f"count in Updated PRISMA tab -> 'Publications' column")            

    def validate_population_col_in_wordreport(self, filepath, locatorname, env):
        self.LogScreenshot.fLogScreenshot(message=f"Validate contents of Population/Sub-group column in Word Report",
                                          pass_=True, log=True, screenshot=False)
        source_template = self.exbase.get_source_template(filepath, 'Sheet1', locatorname)     

        # Read population data values
        pop_list = self.exbase.get_population_data(filepath, 'Sheet1', locatorname)
        # Read slrtype data values
        slrtype = self.exbase.get_slrtype_data(filepath, 'Sheet1', locatorname)         

        self.refreshpage()
        self.go_to_nested_page("importpublications_button", "extraction_upload_btn", env)
        self.imppubpage.upload_file_with_success(locatorname, filepath, env)

        # Go to live slr page
        self.go_to_page("SLR_Homepage", env)
        self.select_data(f"{pop_list[0][0]}", f"{pop_list[0][1]}", env)
        self.select_data(slrtype[0][0], f"{slrtype[0][1]}", env)
        
        self.generate_download_report("word_report", env)
        word_filename = self.get_and_validate_filename(filepath)

        table_count = 5  

        sourcefile = pd.read_excel(f'{source_template[0]}') 
        docs = docx.Document(f'ActualOutputs//{word_filename}')
        try:
            table = docs.tables[table_count]
            data = [[cell.text for cell in row.cells] for row in table.rows]
            df_word = pd.DataFrame(data)

            # Check whether the column name is updated or not
            if 'Population/Sub-group' in df_word.values[0] and 'Population' not in df_word.values[0]:
                self.LogScreenshot.fLogScreenshot(message=f"Column name has been updated from 'Population' to "
                                                          f"'Population/Sub-group'",
                                                  pass_=True, log=True, screenshot=False)
            else:
                self.LogScreenshot.fLogScreenshot(message=f"Column name has not been updated from 'Population' to "
                                                          f"'Population/Sub-group'",
                                                  pass_=False, log=True, screenshot=False)
                raise Exception("Column name has not been updated from 'Population' to 'Population/Sub-group'")

            # Using count variable to loop over columns in word document
            count = 0        
            # df_word.values[0] will give the list of column names from the table in Word document        
            for j in df_word.values[0]:
                # Check the column name is present in Expected test data file
                if j in sourcefile.columns.values:
                    sourcedata = sourcefile[j]
                    word = []
                    for row in docs.tables[table_count].rows:
                        word.append(row.cells[count].text)

                    # Removing NAN/None values if any
                    sourcedata = [item for item in sourcedata if str(item) != 'nan']
                    # Converting Integer list to String list
                    sourcedata = [str(x) for x in sourcedata]
                    # Popping up the column name
                    word.pop(0)

                    comparison_result = self.list_comparison_between_reports_data(sourcedata, word)

                    if len(comparison_result) == 0:
                        self.LogScreenshot.fLogScreenshot(message=f"Content from Table 2-2 is matching with expected "
                                                                  f"test data.",
                                                          pass_=True, log=True, screenshot=False)
                    else:
                        self.LogScreenshot.fLogScreenshot(message=f"Content from Table 2-2 is not matching with "
                                                                  f"expected test data. Mismatch values are "
                                                                  f"arranged in following order -> Expected Excel and "
                                                                  f"Word Report. {comparison_result}",
                                                          pass_=False, log=True, screenshot=False)
                        raise Exception("Elements are not matching between Expected Excel data and Word Report")
                    count += 1
                else:
                    raise Exception("Column names are not matching between Expected Excel data and "
                                    "Word Report")
            self.refreshpage()
            self.go_to_nested_page("importpublications_button", "extraction_upload_btn", env)

            self.imppubpage.delete_file(locatorname, filepath, "file_status_popup_text", "upload_table_rows", env)

            # Go to live slr page
            self.go_to_page("SLR_Homepage", env)
            time.sleep(2)
        except Exception:
            raise Exception("Error in Word report content validation")

    def validate_control_chars_in_wordreport(self, filepath, locatorname, env):
        self.LogScreenshot.fLogScreenshot(message=f"Validate the accessibility of downloaded reports when extraction "
                                                  f"file contains control characters",
                                          pass_=True, log=True, screenshot=False)    

        # Read population data values
        pop_list = self.exbase.get_population_data(filepath, 'Sheet1', locatorname)
        # Read slrtype data values
        slrtype = self.exbase.get_slrtype_data(filepath, 'Sheet1', locatorname)         

        self.refreshpage()
        self.go_to_nested_page("importpublications_button", "extraction_upload_btn", env)
        self.imppubpage.upload_file_with_success(locatorname, filepath, env)

        # Go to live slr page
        self.go_to_page("SLR_Homepage", env)
        try:
            for i in pop_list:
                self.select_data(i[0], i[1], env)
                for j in slrtype:
                    self.select_data(j[0], j[1], env)

                    self.generate_download_report("excel_report", env)
                    excel_filename = self.get_and_validate_filename(filepath)

                    self.generate_download_report("word_report", env)
                    word_filename = self.get_and_validate_filename(filepath)

                    self.preview_result("preview_results", env)
                    self.table_display_check("Table", env)
                    self.generate_download_report("Export_as_excel", env)
                    webexcel_filename = self.get_and_validate_filename(filepath)
                    self.back_to_report_page("Back_to_search_page", env)

                    '''Checking whether we are able to read data from downloaded reports or not'''
                    webexcel = pd.read_excel(f'ActualOutputs//{webexcel_filename}')
                    excel = pd.read_excel(f'ActualOutputs//{excel_filename}')
                    docs = docx.Document(f'ActualOutputs//{word_filename}')

                    # Reading Table 2-2 in word report
                    table = docs.tables[5]
                    data = [[cell.text for cell in row.cells] for row in table.rows]
                    df_word = pd.DataFrame(data)

                    if not webexcel.empty and not excel.empty and not df_word.empty:
                        self.LogScreenshot.fLogScreenshot(message=f"Able to open WebExcel, Complete Excel, Word "
                                                                  f"Reports and read data successfully.",
                                                          pass_=True, log=True, screenshot=False)
                    else:
                        self.LogScreenshot.fLogScreenshot(message=f"Error while opening and reading the data from "
                                                                  f"WebExcel, Complete Excel, Word Reports.",
                                                          pass_=False, log=True, screenshot=False)
                        raise Exception(f"Error while opening and reading the data from WebExcel, Complete Excel, "
                                        f"Word Reports.")
        except Exception:
            raise Exception("Unable to select element")                  

        self.refreshpage()
        self.go_to_nested_page("importpublications_button", "extraction_upload_btn", env)

        self.imppubpage.delete_file(locatorname, filepath, "file_status_popup_text", "upload_table_rows", env)

        # Go to live slr page
        self.go_to_page("SLR_Homepage", env)

    def validate_presence_of_ep_details_in_liveslr_page(self, locatorname, filepath, env):
        self.LogScreenshot.fLogScreenshot(message=f"***Validation of presence of Endpoint Details in LiveSLR page is "
                                                  f"started***", pass_=True, log=True, screenshot=False)

        # Read expected categoris from data sheet
        expected_cats = self.exbase.get_individual_col_data(filepath, locatorname, 'Sheet1', 'Expected_Categories')       

        # Read population data values
        pop_list = self.exbase.get_population_data(filepath, 'Sheet1', locatorname)
        # Read slrtype data values
        slrtype = self.exbase.get_slrtype_data(filepath, 'Sheet1', locatorname)

        # Read extraction file path
        extraction_file = self.exbase.get_template_file_details(filepath, locatorname, 'Files_to_upload')

        # Check the Endpoint details in extraction template
        template_data = openpyxl.load_workbook(f'{extraction_file}')
        template_sheet = template_data['Extraction sheet upload']

        ep_abbr_from_extraction_file = [template_sheet['AO2'].value, template_sheet['BM2'].value,
                                        template_sheet['CR2'].value]

        self.refreshpage()
        self.presence_of_admin_page_option("importpublications_button", env)
        self.go_to_nested_page("importpublications_button", "extraction_upload_btn", env)
        self.imppubpage.upload_file_with_success(locatorname, filepath, env)

        self.go_to_page("SLR_Homepage", env)
        self.select_data(pop_list[0][0], pop_list[0][1], env)
        self.select_data(slrtype[0][0], slrtype[0][1], env)
        
        actual_cats = [i.text for i in self.select_elements('cat_view_eles', env)]      

        for j in ep_abbr_from_extraction_file:
            if j in actual_cats:
                self.LogScreenshot.fLogScreenshot(message=f"Endpoint '{j}' is matching with details present under "
                                                          f"Select Category(ies) to View section in UI",
                                                  pass_=True, log=True, screenshot=True)
            else:
                self.LogScreenshot.fLogScreenshot(message=f"Endpoint '{j}' is not matching with details present under "
                                                          f"Select Category(ies) to View section in UI",
                                                  pass_=False, log=True, screenshot=True)
                raise Exception(f"Endpoint '{j}' is not matching with details present under Select Category(ies) to "
                                f"View section in UI")
        
        cats_comparison = self.exbase.list_comparison_between_reports_data(expected_cats, actual_cats)

        if len(cats_comparison) == 0:
            self.LogScreenshot.fLogScreenshot(message=f"Endpoint details are present in 'Select Category(ies) to "
                                                      f"View' as expected with other filters",
                                              pass_=True, log=True, screenshot=True)
        else:
            self.LogScreenshot.fLogScreenshot(message=f"Mismatch found in Endpoint Details. Mismatch values are "
                                                      f"arranged in following order -> Expected Error Message, "
                                                      f"Actual Error Message. {cats_comparison}",
                                              pass_=False, log=True, screenshot=True)
            raise Exception(f"Mismatch found in Endpont Details.")

        self.go_to_nested_page("importpublications_button", "extraction_upload_btn", env)
        self.imppubpage.delete_file(locatorname, filepath, "file_status_popup_text", "upload_table_rows", env)

        self.LogScreenshot.fLogScreenshot(message=f"***Validation of presence of Endpoint Details in LiveSLR page is "
                                                  f"completed***", pass_=True, log=True, screenshot=False)

    def validate_presence_of_uniquestudies_in_liveslr_page(self, locatorname, filepath, env):
        self.LogScreenshot.fLogScreenshot(message=f"***Validation of presence of Unique Studies for Project Level and "
                                                  f"SLR Type Level in LiveSLR page is started***",
                                          pass_=True, log=True, screenshot=False)

        # Read population data values
        pop_list = self.exbase.get_triple_col_data(filepath, locatorname, 'Sheet1', 'Population',
                                                   'Population_Radio_button', 'Population_stdyno')
        # Read slrtype data values
        slrtype = self.exbase.get_triple_col_data(filepath, locatorname, 'Sheet1', 'slrtype', 'slrtype_Radio_button',
                                                  'slrtype_stdyno')

        # Read extraction file path
        extraction_file = self.exbase.get_template_file_details(filepath, locatorname, 'Files_to_upload')

        extraction_file_data = pd.read_excel(f'{extraction_file}', sheet_name='Extraction sheet upload', skiprows=4)

        self.refreshpage()
        self.presence_of_admin_page_option("importpublications_button", env)
        self.go_to_nested_page("importpublications_button", "extraction_upload_btn", env)
        self.imppubpage.upload_file_with_success(locatorname, filepath, env)

        self.go_to_page("SLR_Homepage", env)       

        for i in pop_list:
            pop_slrstdy_no = self.get_text(i[2], env)
            self.select_data(i[0], i[1], env)            
            for index, j in enumerate(slrtype):
                slr_slrstdy_no = self.get_text(j[2], env)
                self.select_data(j[0], j[1], env)               
        
                # As there is extra data in row number 6, 7, 8 so we are filtering it out based on all the SLR types
                # to get the total Study count at the project level
                filtered_extraction_file_data = extraction_file_data.query('`SLR Type`.str.startswith("Clinical") | '
                                                                           '`SLR Type`.str.startswith("Quality of '
                                                                           'Life") | `SLR Type`.str.startswith('
                                                                           '"Economic") | `SLR Type`.str.startswith('
                                                                           '"Real-world Evidence").values')
                # With the help of above filtered data, we are getting Unique Study count based on Project level
                project_level_stdy_cnt = filtered_extraction_file_data["LiveSLR Study ID"]
                project_level_stdy_cnt_final = []
                # Removing the duplicates
                [project_level_stdy_cnt_final.append(x) for x in list(flatten(project_level_stdy_cnt))
                 if x not in project_level_stdy_cnt_final]
                
                # Get Unique Study ID based on SLR Type
                col_val = extraction_file_data[extraction_file_data["SLR Type"] == j[0]]
                slr_level_stdy_cnt = col_val["LiveSLR Study ID"]
                slr_level_stdy_cnt = [item for item in slr_level_stdy_cnt if str(item) != 'nan']
                slr_level_stdy_cnt_final = []
                # Removing the duplicates
                [slr_level_stdy_cnt_final.append(x) for x in list(flatten(slr_level_stdy_cnt))
                 if x not in slr_level_stdy_cnt_final]

                # Comparison for Unique Study Number
                if int(pop_slrstdy_no) == len(project_level_stdy_cnt_final):
                    self.LogScreenshot.fLogScreenshot(message=f"For '{i[0]}' Project -> Study Number is matching with "
                                                              f"the total number of unique records uploaded from the "
                                                              f"extraction file. Count for '{i[0]}' Project is '"
                                                              f"{len(project_level_stdy_cnt_final)}'",
                                                      pass_=True, log=True, screenshot=False)
                else:
                    self.LogScreenshot.fLogScreenshot(message=f"For '{i[0]}' Project -> Study Number is not matching "
                                                              f"with the total number of unique records uploaded from "
                                                              f"the extraction file. For '{i[0]}' Project -> Count "
                                                              f"displayed in UI is '"
                                                              f"{len(project_level_stdy_cnt_final)}' and Number of "
                                                              f"Unique records uploaded for '{i[0]}' Project is '"
                                                              f"{pop_slrstdy_no}'",
                                                      pass_=False, log=True, screenshot=False)
                    raise Exception(f"For '{i[0]}' Project -> Study Number is not matching with the total number of "
                                    f"unique records uploaded from the extraction file")

                if int(slr_slrstdy_no) == len(slr_level_stdy_cnt_final):
                    self.LogScreenshot.fLogScreenshot(message=f"For '{j[0]}' SLR Type -> Study Number is matching "
                                                              f"with the total number of unique records uploaded from "
                                                              f"the extraction file. Count for '{j[0]}' SLR Type is "
                                                              f"'{len(slr_level_stdy_cnt_final)}'",
                                                      pass_=True, log=True, screenshot=False)
                else:
                    self.LogScreenshot.fLogScreenshot(message=f"For '{j[0]}' SLR Type -> Study Number is not matching "
                                                              f"with the total number of unique records uploaded from "
                                                              f"the extraction file. For '{j[0]}' SLR Type -> Count "
                                                              f"displayed in UI is '{len(slr_level_stdy_cnt_final)}' "
                                                              f"and Number of Unique records uploaded for '{j[0]}' "
                                                              f"SLR Type is '{slr_slrstdy_no}'",
                                                      pass_=False, log=True, screenshot=False)
                    raise Exception(f"For '{j[0]}' SLR Type -> Study Number is not matching with the total number of "
                                    f"unique records uploaded from the extraction file")

        self.go_to_nested_page("importpublications_button", "extraction_upload_btn", env)
        self.imppubpage.delete_file(locatorname, filepath, "file_status_popup_text", "upload_table_rows", env)

        self.go_to_page("SLR_Homepage", env)
        for x in pop_list: 
            # Check absence of Project after deleting the extraction file
            if not self.isvisible(x[0], env, x[0]):
                self.LogScreenshot.fLogScreenshot(message=f"Project '{x[0]}' is not visible in Search LiveSLR page as "
                                                          f"expected after deleting the extraction file.",
                                                  pass_=True, log=True, screenshot=True)
            else:
                self.LogScreenshot.fLogScreenshot(message=f"Project '{x[0]}' is visible in Search LiveSLR page after "
                                                          f"deleting the extraction file.",
                                                  pass_=False, log=True, screenshot=False)
                raise Exception(f"Project '{x[0]}' is visible in Search LiveSLR page")

        self.LogScreenshot.fLogScreenshot(message=f"***Validation of presence of Unique Studies for Project Level and "
                                                  f"SLR Type Level in LiveSLR page is completed***",
                                          pass_=True, log=True, screenshot=False)

    def validate_uniquestudies_reporting_outcomes(self, locatorname, filepath, env):
        self.LogScreenshot.fLogScreenshot(message=f"***Validation of presence of Endpoint Details with Unique Studies "
                                                  f"count in LiveSLR -> Select Studies Reporting Outcome(s) section "
                                                  f"is started***", pass_=True, log=True, screenshot=False)

        # Read population data values
        pop_list = self.exbase.get_population_data(filepath, 'Sheet1', locatorname)
        # Read slrtype data values
        slrtype = self.exbase.get_slrtype_data(filepath, 'Sheet1', locatorname)

        # Read extraction file path
        extraction_file = self.exbase.get_template_file_details(filepath, locatorname, 'Files_to_upload')

        extraction_file_data = pd.read_excel(f'{extraction_file}', sheet_name='Extraction sheet upload', skiprows=3)
        # As there is extra data in row number 6, 7, 8 so we are filtering it out based on all the SLR types
        # to get the total Study count at the project level
        df = extraction_file_data.query('`A-1`.str.startswith("Clinical") | '
                                        '`A-1`.str.startswith("Quality of Life") | '
                                        '`A-1`.str.startswith("Economic") | '
                                        '`A-1`.str.startswith("Real-world Evidence").values')

        self.go_to_page("SLR_Homepage", env)
        self.refreshpage()
        self.presence_of_admin_page_option("importpublications_button", env)
        self.go_to_nested_page("importpublications_button", "extraction_upload_btn", env)
        self.imppubpage.upload_file_with_success(locatorname, filepath, env)

        self.go_to_page("SLR_Homepage", env)
        # self.select_data(pop_list[0][0], pop_list[0][1], env)
        # self.select_data(slrtype[0][0], slrtype[0][1], env)
        for i in pop_list:
            self.select_data(i[0], i[1], env)
            for index, j in enumerate(slrtype):
                self.select_data(j[0], j[1], env)

                self.scroll("reported_variable_section", env)
                self.LogScreenshot.fLogScreenshot(message=f"Values under 'Select Studies Reporting Outcome(s)' "
                                                          f"section : ", pass_=True, log=True, screenshot=True)

                col_id = {'AUT1': ['QA-1', 'QA-2', 'QA-3', 'QB-2', 'QB-6'],
                          'AUT2': ['RC-3', 'RC-4', 'RD-2', 'RD-6'],
                          'AUT3': ['ST-1', 'ST-6', 'ST-10'],
                          'Safety': ['E-3', 'E-4', 'E-5']
                          }

                actual_reported_var = [i.text for i in self.select_elements('reported_variable_section_values', env)]
                actual_reported_var_res = []
                for xy in actual_reported_var:
                    actual_reported_var_res.append(xy.splitlines())

                # Manipulating the col_id dictionary value based on the items displayed in UI.
                res_dict = {}
                for n in actual_reported_var_res:
                    if n[0] in col_id:
                        res_dict[n[0]] = col_id.get(n[0])

                # Above manipulation helps in applying filter for Extraction file to get the unique study count
                index = 0
                for k, v in res_dict.items():
                    res1 = []
                    for item in v:
                        col_val = df[df[item] != "NR"]
                        col_val_res = col_val["A-2"]
                        # Removing the duplicates
                        [res1.append(x) for x in list(flatten(col_val_res)) if x not in res1]

                    # Comparison for Unique Study Number
                    if int(actual_reported_var_res[index][1]) == len(res1):
                        self.LogScreenshot.fLogScreenshot(message=f"For '{i[0]}' Project -> Unique Study Number from "
                                                                  f"'Select Studies Reporting Outcome(s)' section is "
                                                                  f"matching with the total number of unique records "
                                                                  f"uploaded from the extraction file. "
                                                                  f"Count for '{k}' Endpoint is : "
                                                                  f"'{actual_reported_var_res[index][1]}'",
                                                          pass_=True, log=True, screenshot=True)
                    else:
                        self.LogScreenshot.fLogScreenshot(message=f"For '{i[0]}' Project -> Unique Study Number from "
                                                                  f"'Select Studies Reporting Outcome(s)' section is "
                                                                  f"not matching with the total number of unique "
                                                                  f"records uploaded from the extraction file. For "
                                                                  f"'{i[0]}' Project -> Count for '{k}' Endpoint "
                                                                  f"displayed in UI is "
                                                                  f"'{actual_reported_var_res[index][1]}' and "
                                                                  f"Number of Unique records uploaded for "
                                                                  f"'{i[0]}' Project is '{len(res1)}'",
                                                          pass_=False, log=True, screenshot=True)
                        raise Exception(f"For '{i[0]}' Project -> Unique Study Number from 'Select Studies Reporting "
                                        f"Outcome(s)' section is not matching with the total number of unique "
                                        f"records uploaded from the extraction file")
                    index += 1
        
        self.go_to_nested_page("importpublications_button", "extraction_upload_btn", env)
        self.imppubpage.delete_file(locatorname, filepath, "file_status_popup_text", "upload_table_rows", env)

        self.LogScreenshot.fLogScreenshot(message=f"***Validation of presence of Endpoint Details with Unique Studies "
                                                  f"count in LiveSLR -> Select Studies Reporting Outcome(s) section "
                                                  f"is completed***", pass_=True, log=True, screenshot=False)

    def validate_liveslrpage_tooltip(self, locatorname, filepath, env):
        self.LogScreenshot.fLogScreenshot(message=f"***Validation of presence of Tooltip information in Search "
                                                  f"LiveSLR Page is started***",
                                          pass_=True, log=True, screenshot=False)

        expected_tool_tip = self.exbase.get_individual_col_data(filepath, locatorname, 'Sheet1', 'ToolTip_Info')

        actual_tool_tip = []
        tool_tip_locators = {'Select SLR Project': 'slr_pop_tooltip',
                             'Select Type of SLR': 'slr_type_tooltip',
                             'Select Category(ies) to View': 'cat_view_tooltip',
                             'Select Studies Reporting Outcome(s)': 'reporting_outcome_tooltip'}
        
        for k, v in tool_tip_locators.items():
            ele = self.select_element(v, env)
            actions = ActionChains(self.driver)
            actions.move_to_element(ele).perform()
            self.LogScreenshot.fLogScreenshot(message=f"Tool Tip Description for '{k}' is: ",
                                              pass_=True, log=True, screenshot=True)
            actual_tool_tip.append(ele.get_attribute('tooltip'))

        comparison_result = self.list_comparison_between_reports_data(expected_tool_tip, actual_tool_tip)

        if len(comparison_result) == 0:
            self.LogScreenshot.fLogScreenshot(message=f"Tool Tip Description is matching with Expected messages.",
                                              pass_=True, log=True, screenshot=False)
        else:
            self.LogScreenshot.fLogScreenshot(message=f"Tool Tip Description is not matching with Expected messages. "
                                                      f"Mismatch values are arranged in following order -> "
                                                      f"Expected Message and Actual Message. {comparison_result}",
                                              pass_=False, log=True, screenshot=False)
            raise Exception("Tool Tip Description is not matching with Expected messages.")

        self.LogScreenshot.fLogScreenshot(message=f"***Validation of presence of Tooltip information in Search "
                                                  f"LiveSLR Page is completed***",
                                          pass_=True, log=True, screenshot=False)

    # # ############## Using Openpyxl library #################
    # def excel_content_validation(self, webexcel_filename, excel_filename, slrtype):
    #     self.LogScreenshot.fLogScreenshot(message=f"FileNames are: {webexcel_filename} and \n{excel_filename}",
    #                                       pass_=True, log=True, screenshot=False)
    #     sheet_names = ReadConfig.get_sheetname_as_per_slrtype(slrtype)
    #     self.LogScreenshot.fLogScreenshot(message=f"Sheetnames are: {sheet_names}",
    #                                       pass_=True, log=True, screenshot=False)
    #     for sheet in sheet_names:
    #         webex = []
    #         compex = []
    #         webexcel = openpyxl.load_workbook(f'ActualOutputs//{webexcel_filename}')
    #         webexcel_sheet = webexcel[sheet]
    #         excel = openpyxl.load_workbook(f'ActualOutputs//{excel_filename}')
    #         excel_sheet = excel[sheet]
    #
    #         try:
    #             for i in webexcel_sheet.rows:
    #                 webex.append(i[0].value)
    #
    #             for j in excel_sheet.rows:
    #                 compex.append(j[0].value)
    #
    #             # Removing None values from the lists
    #             webex = list(filter(None, webex))
    #             compex = list(filter(None, compex))
    #
    #             # Comparison and Merging with regular expressions
    #             res = [ele for ele in compex if (ele in webex)]
    #             self.LogScreenshot.fLogScreenshot(message=f"Regular Expression output: {res}",
    #                                               pass_=True, log=True, screenshot=False)
    #             if str(bool(res)):
    #                 # Extracting numbers from the lists
    #                 ele1 = [x for x in webex if isinstance(x, numbers.Number)]
    #                 ele2 = [x for x in compex if isinstance(x, numbers.Number)]
    #                 if ele1 == sorted(ele1) and ele2 == sorted(ele2):
    #                     if ele1 == ele2:
    #                         self.LogScreenshot.fLogScreenshot(message=f"Study Identifier are matching between "
    #                                                                   f"WebExcel and Complete Excel report. \n"
    #                                                                   f"Compared Element values are: {ele1} "
    #                                                                   f"and \n{ele2}",
    #                                                           pass_=True, log=True, screenshot=False)
    #                     else:
    #                         self.LogScreenshot.fLogScreenshot(message=f"Study Identifier are not matching between "
    #                                                                   f"WebExcel and Complete Excel report. \n"
    #                                                                   f"Compared Element values are: {ele1} "
    #                                                                   f"and \n{ele2}",
    #                                                           pass_=False, log=True, screenshot=False)
    #                         raise Exception("Study Identifier are not matching")
    #                 else:
    #                     raise Exception("Study Identifier Elements are not in sorted order")
    #         except Exception:
    #             raise Exception("Error in Excel sheet content validation")
    #
    # def excel_to_word_content_validation(self, webexcel_filename, excel_filename, word_filename, slrtype):
    #     self.LogScreenshot.fLogScreenshot(message=f"FileNames are: "
    #                                               f"{webexcel_filename}, \n{excel_filename}, \n{word_filename}",
    #                                       pass_=True, log=True, screenshot=False)
    #     # Index of Table number 6 is : 5. Starting point for word table content comparison
    #     table_count = 5
    #     sheet_names = ReadConfig.get_sheetname_as_per_slrtype(slrtype)
    #     self.LogScreenshot.fLogScreenshot(message=f"Sheetnames are: {sheet_names}",
    #                                       pass_=True, log=True, screenshot=False)
    #     for sheet in sheet_names:
    #         webex = []
    #         compex = []
    #         word = []
    #         webexcel = openpyxl.load_workbook(f'ActualOutputs//{webexcel_filename}')
    #         webexcel_sheet = webexcel[sheet]
    #         excel = openpyxl.load_workbook(f'ActualOutputs//{excel_filename}')
    #         excel_sheet = excel[sheet]
    #         docs = docx.Document(f'ActualOutputs//{word_filename}')
    #         try:
    #             for i in webexcel_sheet.rows:
    #                 webex.append(i[3].value)
    #
    #             for j in excel_sheet.rows:
    #                 compex.append(j[3].value)
    #
    #             for row in docs.tables[table_count].rows:
    #                 word.append(row.cells[0].text)
    #             # Incrementing the counter to switch for next table based on the number of excel sheet names
    #             table_count += 1
    #
    #             # Removing None values from the lists
    #             webex = list(filter(None, webex))
    #             compex = list(filter(None, compex))
    #
    #             if compex == webex and compex == word:
    #                 self.LogScreenshot.fLogScreenshot(message=f"Contents are matching in WebExcel, Complete "
    #                                                           f"Excel and Complete Word Reports. "
    #                                                           f"Short Reference values are: \n"
    #                                                           f"{webex} \n {compex} \n {word}",
    #                                                   pass_=True, log=True, screenshot=False)
    #             else:
    #                 self.LogScreenshot.fLogScreenshot(message=f"Contents are not matching in WebExcel, Complete "
    #                                                           f"Excel and Complete Word Reports. "
    #                                                           f"Short Reference values are: \n"
    #                                                           f"{webex} \n {compex} \n {word}",
    #                                                   pass_=False, log=True, screenshot=False)
    #                 raise Exception("Word report contents are not matching")
    #         except Exception:
    #             raise Exception("Error in Word report content validation")
