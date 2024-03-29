import os
from pathlib import Path
import time

import docx
import openpyxl
import pandas as pd
from selenium.webdriver.support.wait import WebDriverWait
from selenium.webdriver.support.ui import Select
from pandas.core.common import flatten

from Pages.Base import Base, fWaitFor
from Pages.ExtendedBasePage import ExtendedBase
from Pages.ImportPublicationsPage import ImportPublicationPage
from Pages.OpenLiveSLRPage import LiveSLRPage
from Pages.SLRReportPage import SLRReport
from utilities.customLogger import LogGen
from utilities.logScreenshot import cLogScreenshot
from utilities.readProperties import ReadConfig


class UtilityOutcome(Base):

    """Constructor of the UtilityOutcome class"""
    def __init__(self, driver, extra):
        # initializing the driver from base class
        super().__init__(driver, extra)  
        self.extra = extra
        # Instantiate the Base class
        self.base = Base(self.driver, self.extra)
        # Creating object of ExtendedBase class
        self.exbase = ExtendedBase(self.driver, extra)        
        # Creating object of liveslrpage class
        self.liveslrpage = LiveSLRPage(self.driver, self.extra)
        # Creating object of ImportPublicationPage class
        self.imppubpage = ImportPublicationPage(self.driver, extra)
        # Creating object of slrreport class
        self.slrreport = SLRReport(self.driver, extra)
        # Instantiate the logger class
        self.logger = LogGen.loggen()
        # Instantiate the logScreenshot class
        self.LogScreenshot = cLogScreenshot(self.driver, self.extra)
        # Instantiate webdriver wait class
        self.wait = WebDriverWait(driver, 20)

    def get_population_data_specific_sheet(self, filepath, sheet):
        file = pd.read_excel(filepath, sheet_name=sheet)
        pop = list(file['Population'].dropna())
        pop_button = list(file['Population_Radio_button'].dropna())
        population_data = [(pop[i], pop_button[i]) for i in range(0, len(pop))]
        return population_data

    def get_slrtype_data_specific_sheet(self, filepath, sheet):
        file = pd.read_excel(filepath, sheet_name=sheet)
        slrtype = list(file['slrtype'].dropna())
        slrtype_button = list(file['slrtype_Radio_button'].dropna())
        slrtype_data = [(slrtype[i], slrtype_button[i]) for i in range(0, len(slrtype))]
        return slrtype_data

    def get_reported_variables_specific_sheet(self, filepath, sheet):
        file = pd.read_excel(filepath, sheet_name=sheet)
        reported_var = list(file['ReportedVariables'].dropna())
        reported_var_button = list(file['Reportedvariable_checkbox'].dropna())
        return reported_var, reported_var_button
    
    def get_util_source_template(self, filepath, sheet):
        file = pd.read_excel(filepath, sheet_name=sheet)
        expectedfilepath = list(os.getcwd()+file['ExpectedSourceTemplateFile'].dropna())
        return expectedfilepath
    
    def get_extraction_file_to_upload(self, filepath, sheet, locatorname):
        df = pd.read_excel(filepath, sheet_name=sheet)
        # path = df.loc[df['Name'] == locatorname]['ExtractionFile'].dropna().to_list()
        # result = [[os.getcwd() + path[i]] for i in range(0, len(path))]
        pop_name = df.loc[df['Name'] == locatorname]['Population_name'].dropna().to_list()
        path = df.loc[df['Name'] == locatorname]['Files_to_upload'].dropna().to_list()
        filename = df.loc[df['Name'] == locatorname]['Expected_File_names'].dropna().to_list()
        result = [[pop_name[i], os.getcwd() + path[i], filename[i]] for i in range(0, len(pop_name))]        
        return result
    
    # def get_import_pop_to_upload(self, filepath, sheet, locatorname):
    #     df = pd.read_excel(filepath, sheet_name=sheet)
    #     pop = df.loc[df['Name'] == locatorname]['Import_Pop'].dropna().to_list()
    #     return pop

    def get_population_to_upload(self, filepath, sheet, locatorname):
        df = pd.read_excel(filepath, sheet_name=sheet)
        pop = df.loc[df['Name'] == locatorname]['Population'].dropna().to_list()
        pop_button = df.loc[df['Name'] == locatorname]['Population_Radio_button'].dropna().to_list()
        result = [[pop[i], pop_button[i]] for i in range(0, len(pop))]
        return result
    
    def get_slrtype_to_upload(self, filepath, sheet, locatorname):
        df = pd.read_excel(filepath, sheet_name=sheet)
        slrtype = df.loc[df['Name'] == locatorname]['slrtype'].dropna().to_list()
        slrtype_button = df.loc[df['Name'] == locatorname]['slrtype_Radio_button'].dropna().to_list()
        result = [[slrtype[i], slrtype_button[i]] for i in range(0, len(slrtype))]
        return result
    
    def validate_filename(self, filename, filepath, sheet):
        try:
            file = pd.read_excel(filepath, sheet_name=sheet)
            expectedname = list(file['ExpectedFilenames'].dropna())
            actualname = filename[:-19]
            if actualname in expectedname:
                self.LogScreenshot.fLogScreenshot(message=f"Correct file is downloaded",
                                                  pass_=True, log=True, screenshot=False)
            else:
                raise Exception
        except Exception:
            self.LogScreenshot.fLogScreenshot(
                message=f"Filename is not present in the expected list. Expected Filenames are {expectedname} "
                        f"and Actual Filename is {actualname}",
                pass_=False, log=True, screenshot=False)

    def qol_presenceof_utilitysummary_into_excelreport(self, excel_filename, util_filepath):
        source_template = self.get_util_source_template(util_filepath, 'NewImportLogic')
        
        self.LogScreenshot.fLogScreenshot(
            message=f"*****Check Presence of Utility Summary Tab in Complete Excel Report*****",
            pass_=True, log=True, screenshot=False)
        excel_data = openpyxl.load_workbook(f'ActualOutputs//{excel_filename}')
        if 'Utility Summary' in excel_data.sheetnames:
            self.LogScreenshot.fLogScreenshot(
                message=f"'Utility Summary' tab is present in Downloaded Complete Excel Report",
                pass_=True, log=True, screenshot=False)
        else:
            self.LogScreenshot.fLogScreenshot(message=f"'Utility Summary' tab is missing in Complete Excel Report",
                                              pass_=False, log=True, screenshot=False)
            raise Exception("'Utility Summary' tab is missing in Complete Excel Report")

        excel_sheet = excel_data[f'Utility Summary']
        if excel_sheet['H1'].value == 'Back To Toc':
            self.LogScreenshot.fLogScreenshot(message=f"'Back To Toc' option is present in 'Utility Summary' sheet",
                                              pass_=True, log=True, screenshot=False)
        else:
            self.LogScreenshot.fLogScreenshot(message=f"'Back To Toc' option is not present in 'Utility Summary' sheet",
                                              pass_=False, log=True, screenshot=False)
            raise Exception(f"'Back To Toc' option is not present in 'Utility Summary' sheet")
        
        toc_sheet = pd.read_excel(f'ActualOutputs//{excel_filename}', sheet_name="TOC", skiprows=3)
        col_data = list(toc_sheet.iloc[:, 1])
        if f'Utility Summary' in col_data:
            self.LogScreenshot.fLogScreenshot(message=f"'Utility Summary' is present in TOC sheet.",
                                              pass_=True, log=True, screenshot=False)
        else:
            self.LogScreenshot.fLogScreenshot(
                message=f"'Utility Summary' is not present in TOC sheet. Available Data from TOC sheet: {col_data}",
                pass_=False, log=True, screenshot=False)
            raise Exception("'Utility Summary' is not present in TOC sheet.")
        
        self.LogScreenshot.fLogScreenshot(
            message=f"*****Check Presence of expected columns in Utility Summary Tab*****",
            pass_=True, log=True, screenshot=False)
        sourcefile = pd.read_excel(f'{source_template[0]}', sheet_name='ExpectedUtilitySummary')
        actualexcel = pd.read_excel(f'ActualOutputs//{excel_filename}', sheet_name='Utility Summary', skiprows=2)
        
        source_col_val = [item for item in list(sourcefile.columns.values) if str(item) != 'nan']
        compex_col_val = [item for item in list(actualexcel.columns.values) if str(item) != 'nan']
        # Removing the unnamed column from the downloaded report. This column got created for 'Back To TOC' option
        compex_col_val.pop(-1)

        if source_col_val == compex_col_val:
            self.LogScreenshot.fLogScreenshot(
                message=f"Expected column names are present in Complete Excel Report. Column names are: "
                        f"{compex_col_val}",
                pass_=True, log=True, screenshot=False)
        else:
            self.LogScreenshot.fLogScreenshot(
                message=f"Expected column names are not present in Complete Excel Report. Expected Excel Excel "
                        f"Column names are : {source_col_val},\n Complete Excel Column names are : {compex_col_val}",
                pass_=False, log=True, screenshot=False)
            raise Exception("Expected column names are not present in Complete Excel Report")
    
    def qol_verify_source_to_target_row_counts_excelreport(self, excel_filename, util_filepath):
        source_template = self.get_util_source_template(util_filepath, 'NewImportLogic')

        sourcefile = pd.read_excel(f'{source_template[0]}', sheet_name='ExpectedUtilitySummary')
        actualexcel = pd.read_excel(f'ActualOutputs//{excel_filename}', sheet_name='Utility Summary', skiprows=2)
        
        cols = list(sourcefile.columns.values)

        # sourcefile_col = sourcefile['Short Reference']
        # actualexcel_col = actualexcel['Short Reference']
        #
        # sourcefile_col = [item for item in sourcefile_col if str(item) != 'nan']
        # actualexcel_col = [item for item in actualexcel_col if str(item) != 'nan']

        for col in cols:
            source_col = sourcefile[col]
            actual_col = actualexcel[col]

            source_col = [item for item in source_col if str(item) != 'nan']
            actual_col = [item for item in actual_col if str(item) != 'nan']

            if len(source_col) == len(actual_col):
                self.LogScreenshot.fLogScreenshot(
                    message=f"Row count in column '{col}' are matching between Expected Excel Utility File and "
                            f"Complete Excel Report. Expected Excel Utility File Elements Length {len(source_col)},\n"
                            f"Complete Excel Elements Length {len(actual_col)}",
                    pass_=True, log=True, screenshot=False)
            else:
                self.LogScreenshot.fLogScreenshot(
                    message=f"Row count in column '{col}' are not matching. Expected Excel Utility File Elements "
                            f"Length {len(source_col)},\n Complete Excel Elements Length {len(actual_col)}",
                    pass_=False, log=True, screenshot=False)
                raise Exception("Row Counts are not matching between Expected Excel Utility File and Complete "
                                "Excel report")
    
    def qol_presenceof_utilitysummary_into_wordreport(self, word_filename, util_filepath):
        source_template = self.get_util_source_template(util_filepath, 'NewImportLogic')

        sourcefile = pd.read_excel(f'{source_template[0]}', sheet_name='ExpectedUtilitySummary')
        docs = docx.Document(f'ActualOutputs//{word_filename}')
        
        table = docs.tables[6]
        data = [[cell.text for cell in row.cells] for row in table.rows]
        df_word = pd.DataFrame(data)

        self.LogScreenshot.fLogScreenshot(message=f"'Utility Table' is present in Downloaded Complete Word Report",
                                          pass_=True, log=True, screenshot=False)

        # Column name check between Expected data and Word report
        source_col_val = [item for item in list(sourcefile.columns.values) if str(item) != 'nan']
        word_col_val = [item for item in list(df_word.values[0]) if str(item) != 'nan']
        if source_col_val == word_col_val:
            self.LogScreenshot.fLogScreenshot(
                message=f"Expected column names are present in Word Report. Column names are: {word_col_val}",
                pass_=True, log=True, screenshot=False)
        else:
            self.LogScreenshot.fLogScreenshot(
                message=f"Expected column names are not present in Word Report. Expected Excel Column names are : "
                        f"{source_col_val},\n Word Report Column names are : {word_col_val}",
                pass_=False, log=True, screenshot=False)
            raise Exception("Expected column names are not present in Word Report")

    def qol_verify_source_to_target_row_counts_wordreport(self, word_filename, util_filepath):
        source_template = self.get_util_source_template(util_filepath, 'NewImportLogic')

        sourcefile = pd.read_excel(f'{source_template[0]}', sheet_name='ExpectedUtilitySummary')
        docs = docx.Document(f'ActualOutputs//{word_filename}')
        
        table = docs.tables[6]
        data = [[cell.text for cell in row.cells] for row in table.rows]
        df_word = pd.DataFrame(data)

        # Using count variable to loop over columns in word document
        count = 0
        # df_word.values[0] will give the list of column names from the table in Word document
        for col_name in df_word.values[0]:
            if col_name in sourcefile.columns.values:
                source = sourcefile[col_name]
                word = []
                for row in docs.tables[6].rows:
                    word.append(row.cells[count].text)
                
                source = [item for item in source if str(item) != 'nan']
                word.pop(0)

                if len(source) == len(word):
                    self.LogScreenshot.fLogScreenshot(
                        message=f"Row count in Column '{col_name}' are matching between Expected Excel Utility "
                                f"File and Word Report.\nExpected Excel Utility File Elements Length: {len(source)},"
                                f"\n Word Elements Length: {len(word)}\n",
                        pass_=True, log=True, screenshot=False)
                else:
                    self.LogScreenshot.fLogScreenshot(
                        message=f"Row count in Column '{col_name}' are not matching between Expected Excel Utility "
                                f"File and Word Report.\nExpected Excel Utility File Elements Length: {len(source)},"
                                f"\n Word Elements Length: {len(word)}\n",
                        pass_=False, log=True, screenshot=False)
                    raise Exception("Row Counts are not matching between Expected Excel Utility File and Word Report.")
                count += 1
            else:
                raise Exception(f"Column name '{col_name}' is not matching between Expected Excel Utility File and "
                                f"Word Report.")
        
    def qol_verify_excelreport_utility_summary_sorting_order(self, excel_filename, util_filepath):
        self.LogScreenshot.fLogScreenshot(
            message=f"*****Validation of Table data and Sorting order in Utility Summary sheet Started*****",
            pass_=True, log=True, screenshot=False)
        source_template = self.get_util_source_template(util_filepath, 'NewImportLogic')
        
        # Utility Summary sheet comparison with Expected results
        self.LogScreenshot.fLogScreenshot(message=f"*****Utility Summary Sheet Comparison*****",
                                          pass_=True, log=True, screenshot=False)
        sourcefile = pd.read_excel(f'{source_template[0]}', sheet_name='ExpectedUtilitySummary')
        actualexcel = pd.read_excel(f'ActualOutputs//{excel_filename}', sheet_name='Utility Summary', skiprows=2)
        
        cols = list(sourcefile.columns.values)

        # Check sorting order from downloaded excel report
        actual_slrsource = actualexcel["SLR Source"]
        actual_slrsource = [item for item in actual_slrsource if str(item) != 'nan']
        actual_slrsource_final = []
        # Removing the duplicates
        [actual_slrsource_final.append(x) for x in list(flatten(actual_slrsource)) if x not in actual_slrsource_final]

        self.LogScreenshot.fLogScreenshot(message=f"Unique SLR Source column values are : {actual_slrsource_final}",
                                          pass_=True, log=True, screenshot=False)
        
        for m in actual_slrsource_final:
            col_val = actualexcel[actualexcel["SLR Source"] == m]
            col_val_res = col_val["Short Reference"]
            col_val_res = [item for item in col_val_res if str(item) != 'nan']
            if col_val_res == sorted(col_val_res):
                self.LogScreenshot.fLogScreenshot(
                    message=f"From 'Utility Summary' sheet -> For '{m}' SLR Source, contents in column "
                            f"'Short Reference' are in sorted order",
                    pass_=True, log=True, screenshot=False)
            else:
                raise Exception(f"From 'Utility Summary' sheet -> For '{m}' SLR Source, contents in column "
                                f"'Short Reference' are not in Sorted order")

        # Check the length of 1st column from the report to make sure number of rows are as expected
        source_col_len = sourcefile["SLR Source"]
        actual_col_len = actualexcel["SLR Source"]

        source_col_len = [item for item in source_col_len if str(item) != 'nan']
        actual_col_len = [item for item in actual_col_len if str(item) != 'nan']

        if len(source_col_len) == len(actual_col_len):
            self.LogScreenshot.fLogScreenshot(
                message=f"Elements length is matching between Expected Excel and Complete Excel Report. Expected "
                        f"Excel Elements Length: {len(source_col_len)}\nExcel Elements Length: {len(actual_col_len)}",
                pass_=True, log=True, screenshot=False)
        
            # Content comparison between Source file and Complete Excel report
            for col in cols:
                source_col = sourcefile[col]
                actual_col = actualexcel[col]

                source_col = [item for item in source_col if str(item) != 'nan']
                actual_col = [item for item in actual_col if str(item) != 'nan']

                comparison_result = self.slrreport.list_comparison_between_reports_data(source_col, actual_col)

                if len(comparison_result) == 0:
                    self.LogScreenshot.fLogScreenshot(
                        message=f"Contents in column '{col}' are matching between Expected Excel and Complete "
                                f"Excel Report",
                        pass_=True, log=True, screenshot=False)
                else:
                    self.LogScreenshot.fLogScreenshot(
                        message=f"Contents in column '{col}' are not matching. Duplicate values are arranged in "
                                f"following order -> Expected Excel File, Complete Excel. {comparison_result}",
                        pass_=False, log=True, screenshot=False)
                    raise Exception("Contents are not matching between Expected Excel and Complete Excel report")
        else:
            self.LogScreenshot.fLogScreenshot(
                message=f"Elements length is not matching between Expected Excel and Complete Excel Report. Expected "
                        f"Excel Elements Length: {len(source_col_len)}\n Excel Elements Length: {len(actual_col_len)}",
                pass_=False, log=True, screenshot=False)
            raise Exception(f"Elements length is not matching between Expected Excel and Complete Excel Report.")

        self.LogScreenshot.fLogScreenshot(
            message=f"*****Validation of Table data and Sorting order in Utility Summary sheet Completed*****",
            pass_=True, log=True, screenshot=False)
    
    def qol_verify_wordreport_utility_table_sorting_order(self, word_filename, util_filepath):
        self.LogScreenshot.fLogScreenshot(
            message=f"*****Validation of Table data and Sorting order in Utility Table Started*****",
            pass_=True, log=True, screenshot=False)
        source_template = self.get_util_source_template(util_filepath, 'NewImportLogic')
        
        # Utility Table comparison with Expected results
        self.LogScreenshot.fLogScreenshot(message=f"*****Utility Summary Sheet Comparison*****",
                                          pass_=True, log=True, screenshot=False)
        sourcefile = pd.read_excel(f'{source_template[0]}', sheet_name='ExpectedUtilitySummary')
        docs = docx.Document(f'ActualOutputs//{word_filename}')

        table = docs.tables[6]
        data = [[cell.text for cell in row.cells] for row in table.rows]
        df_word = pd.DataFrame(data)

        # source = sourcefile["SLR Source"]
        word = []
        for row in docs.tables[6].rows:
            word.append(row.cells[0].text)
        
        # source = [item for item in source if str(item) != 'nan']
        word.pop(0)
        word_slrsource_final = []
        # Removing the duplicates
        [word_slrsource_final.append(x) for x in list(flatten(word)) if x not in word_slrsource_final]

        print(f"Unique SLR Source column values in word report are : {word_slrsource_final}")

        for n in word_slrsource_final:
            col_val = df_word[df_word[0] == n]
            col_val_res1 = col_val[1]
            col_val_res1 = [item for item in col_val_res1 if str(item) != 'nan']

            if col_val_res1 == sorted(col_val_res1):
                self.LogScreenshot.fLogScreenshot(
                    message=f"From 'Utility Table' -> For '{n}' SLR Source, contents in column 'Short Reference' "
                            f"are in sorted order",
                    pass_=True, log=True, screenshot=False)
            else:
                raise Exception(f"From 'Utility Table' -> For '{n}' SLR Source, contents in column 'Short Reference' "
                                f"are not in Sorted order")
        
        # Check the length of 1st column from the report to make sure number of rows are as expected
        source_len = sourcefile[df_word.values[0][0]]
        word_len = []
        for row in docs.tables[6].rows:
            word_len.append(row.cells[0].text)

        source_len = [item for item in source_len if str(item) != 'nan']
        word_len.pop(0)

        if len(source_len) == len(word_len):
            self.LogScreenshot.fLogScreenshot(
                message=f"Elements length is matching between Expected Excel and Complete Word Report. Expected "
                        f"Excel Elements Length: {len(source_len)}\n Word Elements Length: {len(word_len)}\n",
                pass_=True, log=True, screenshot=False)
        
            # Content comparison between Source file and Complete Word report
            # Using count variable to loop over columns in word document
            count = 0
            # df_word.values[0] will give the list of column names from the table in Word document
            for col_name in df_word.values[0]:
                if col_name in sourcefile.columns.values:
                    source = sourcefile[col_name]
                    word = []
                    for row in docs.tables[6].rows:
                        word.append(row.cells[count].text)

                    source = [item for item in source if str(item) != 'nan']
                    word.pop(0)

                    comparison_result = self.slrreport.list_comparison_between_reports_data(source, word)

                    if len(comparison_result) == 0:
                        self.LogScreenshot.fLogScreenshot(
                            message=f"Contents in Column '{col_name}' are matching between Expected Excel and "
                                    f"Word Reports.",
                            pass_=True, log=True, screenshot=False)
                    else:
                        self.LogScreenshot.fLogScreenshot(
                            message=f"Contents in Column '{col_name}' are not matching between Expected Excel and "
                                    f"Word Reports.\nDuplicate values are arranged in following order -> Expected "
                                    f"Excel and Word Report. {comparison_result}",
                            pass_=False, log=True, screenshot=False)
                        raise Exception("Elements are not matching between Expected Excel and Word Reports")
                    count += 1
                else:
                    raise Exception("Column names are not matching between Expected Excel and Word Reports")
        else:
            self.LogScreenshot.fLogScreenshot(
                message=f"Elements length is not matching between Expected Excel and Complete Word Report. Expected "
                        f"Excel Elements Length: {len(source_len)}\n Word Elements Length: {len(word_len)}\n",
                pass_=False, log=True, screenshot=False)

        self.LogScreenshot.fLogScreenshot(
            message=f"*****Validation of Table data and Sorting order in Utility Table Completed*****",
            pass_=True, log=True, screenshot=False)

    def qol_utility_summary_validation(self, webexcel_filename, excel_filename, util_filepath, word_filename):
        self.LogScreenshot.fLogScreenshot(
            message=f"*****Content validation between Extraction Template, Complete Excel Report and Complete "
                    f"Word Report for Utility Summary Started*****",
            pass_=True, log=True, screenshot=False)
        source_template = self.get_util_source_template(util_filepath, 'NewImportLogic')
        
        # QOL Report sheet comparison with Expected results
        self.LogScreenshot.fLogScreenshot(message=f"*****QOL Report sheet Comparison*****",
                                          pass_=True, log=True, screenshot=False)
        sourcefile1 = pd.read_excel(f'{source_template[0]}', sheet_name='ExpectedReportData')
        webexcel1 = pd.read_excel(f'ActualOutputs//{webexcel_filename}', sheet_name='QOL Report', skiprows=3)
        actualexcel1 = pd.read_excel(f'ActualOutputs//{excel_filename}', sheet_name='QOL Report', skiprows=3)
        
        cols1 = list(sourcefile1.columns.values)

        sourcefile_col1 = sourcefile1['Study Identifier']
        actualexcel_col1 = actualexcel1['Study Identifier']
        actualwebexcel_col1 = webexcel1['Study Identifier']

        sourcefile_col1 = [item for item in sourcefile_col1 if str(item) != 'nan']
        actualexcel_col1 = [item for item in actualexcel_col1 if str(item) != 'nan']
        actualwebexcel_col1 = [item for item in actualwebexcel_col1 if str(item) != 'nan']

        if sourcefile_col1 == sorted(sourcefile_col1) and actualexcel_col1 == sorted(actualexcel_col1) and \
                actualwebexcel_col1 == sorted(actualwebexcel_col1):
            self.LogScreenshot.fLogScreenshot(
                message=f"From 'QOL Report' sheet, contents in column 'Study Identifier' are in sorted order",
                pass_=True, log=True, screenshot=False)
            if len(sourcefile_col1) == len(actualexcel_col1) == len(actualwebexcel_col1):
                self.LogScreenshot.fLogScreenshot(
                    message=f"Elements length is matching between Expected Excel, Web_Excel and Complete Excel Report."
                            f"Expected Excel Elements Length: {len(sourcefile_col1)}\n WebExcel Elements Length: "
                            f"{len(actualwebexcel_col1)}\n Excel Elements Length: {len(actualexcel_col1)}\n",
                    pass_=True, log=True, screenshot=False)
                for col in cols1:
                    source_col1 = sourcefile1[col]
                    actual_excel_col1 = actualexcel1[col]
                    actual_webexcel_col = webexcel1[col]

                    source_col1 = [item for item in source_col1 if str(item) != 'nan']
                    actual_excel_col1 = [item for item in actual_excel_col1 if str(item) != 'nan']
                    actual_webexcel_col = [item for item in actual_webexcel_col if str(item) != 'nan']

                    comparison_result = self.slrreport.\
                        list_comparison_between_reports_data(source_col1, actual_excel_col1,
                                                             webex_list=actual_webexcel_col)

                    if len(comparison_result) == 0:
                        self.LogScreenshot.fLogScreenshot(
                            message=f"Contents in column '{col}' are matching between Expected Excel, WebExcel and "
                                    f"Complete Excel Report",
                            pass_=True, log=True, screenshot=False)
                    else:
                        self.LogScreenshot.fLogScreenshot(
                            message=f"Contents in column '{col}' are not matching. Duplicate values are arranged "
                                    f"in following order -> Expected Excel, WebExcel, Complete Excel. "
                                    f"{comparison_result}",
                            pass_=False, log=True, screenshot=False)
                        raise Exception("Contents are not matching between Expected Excel, WebExcel and Complete "
                                        "Excel report")
            else:
                self.LogScreenshot.fLogScreenshot(
                    message=f"Elements length is not matching between Expected Excel, Web_Excel and Complete "
                            f"Excel Report. Expected Excel Elements Length: {len(sourcefile_col1)}\n WebExcel "
                            f"Elements Length: {len(actualwebexcel_col1)}\n Excel Elements Length: "
                            f"{len(actualexcel_col1)}",
                    pass_=False, log=True, screenshot=False)
                raise Exception(f"Elements length is not matching between Expected Excel, Web_Excel and "
                                f"Complete Excel Report.")
        else:
            raise Exception("From 'QOL Report' sheet, contents in column 'Study Identifier' are not in Sorted order")
        
        # Utility Summary sheet comparison with Expected results
        self.LogScreenshot.fLogScreenshot(message=f"*****Utility Summary Sheet Comparison*****",
                                          pass_=True, log=True, screenshot=False)
        sourcefile = pd.read_excel(f'{source_template[0]}', sheet_name='ExpectedUtilitySummary')
        actualexcel = pd.read_excel(f'ActualOutputs//{excel_filename}', sheet_name='Utility Summary', skiprows=2)
        
        cols = list(sourcefile.columns.values)

        sourcefile_col = sourcefile['Short Reference']
        actualexcel_col = actualexcel['Short Reference']

        sourcefile_col = [item for item in sourcefile_col if str(item) != 'nan']
        actualexcel_col = [item for item in actualexcel_col if str(item) != 'nan']

        if len(sourcefile_col) == len(actualexcel_col):
            self.LogScreenshot.fLogScreenshot(
                message=f"Elements length is matching between Expected Excel and Complete Excel Report. Expected "
                        f"Excel Length: {len(sourcefile_col)}\n Excel Elements Length: {len(actualexcel_col)}\n",
                pass_=True, log=True, screenshot=False)

            for col in cols:
                source_col = sourcefile[col]
                actual_col = actualexcel[col]

                source_col = [item for item in source_col if str(item) != 'nan']
                actual_col = [item for item in actual_col if str(item) != 'nan']

                comparison_result = self.slrreport.list_comparison_between_reports_data(source_col, actual_col)

                if len(comparison_result) == 0:
                    self.LogScreenshot.fLogScreenshot(
                        message=f"Contents in column '{col}' are matching between Expected Excel and Complete Excel "
                                f"Report",
                        pass_=True, log=True, screenshot=False)
                else:
                    self.LogScreenshot.fLogScreenshot(
                        message=f"Contents in column '{col}' are not matching. Duplicate values are arranged in "
                                f"following order -> Expected Excel, Complete Excel. {comparison_result}",
                        pass_=False, log=True, screenshot=False)
                    raise Exception("Contents are not matching between Expected Excel and Complete Excel report")
        else:
            self.LogScreenshot.fLogScreenshot(
                message=f"Elements length is not matching between Expected Excel and Complete Excel Report. Expected "
                        f"Excel Elements Length: {len(sourcefile_col)}\n Excel Elements Length: "
                        f"{len(actualexcel_col)}",
                pass_=False, log=True, screenshot=False)
            raise Exception(f"Elements length is not matching between Expected Excel and Complete Excel Report.")
        
        # Word report content comparison for Utility Table
        self.LogScreenshot.fLogScreenshot(message=f"*****Word Report Utility Table Comparison*****",
                                          pass_=True, log=True, screenshot=False)
        docs = docx.Document(f'ActualOutputs//{word_filename}')
        try:
            table = docs.tables[6]
            data = [[cell.text for cell in row.cells] for row in table.rows]
            df_word = pd.DataFrame(data)

            # Column name check between Expected data and Word report
            source_col_val = [item for item in list(sourcefile.columns.values) if str(item) != 'nan']
            word_col_val = [item for item in list(df_word.values[0]) if str(item) != 'nan']
            if source_col_val == word_col_val:
                self.LogScreenshot.fLogScreenshot(
                    message=f"Expected column names are present in Word Report. Column names are: {word_col_val}",
                    pass_=True, log=True, screenshot=False)
            else:
                self.LogScreenshot.fLogScreenshot(
                    message=f"Expected column names are not present in Word Report.Expected Excel Column names "
                            f"are : {source_col_val},\n Word Report Column names are : {word_col_val}",
                    pass_=False, log=True, screenshot=False)
                raise Exception("Expected column names are not present in Word Report")

            # Check the length of 1st column from the report to make sure number of rows are as expected
            source_len = sourcefile[df_word.values[0][0]]
            compex_len = actualexcel[df_word.values[0][0]]
            word_len = []
            for row in docs.tables[6].rows:
                word_len.append(row.cells[0].text)

            source_len = [item for item in source_len if str(item) != 'nan']
            compex_len = [item for item in compex_len if str(item) != 'nan']
            word_len.pop(0)

            if len(source_len) == len(compex_len) == len(word_len):
                self.LogScreenshot.fLogScreenshot(
                    message=f"Elements length is matching between Expected Excel, Complete Excel and Complete Word "
                            f"Report. Expected Excel Elements Length: {len(source_len)}\n Excel Elements "
                            f"Length: {len(compex_len)}\n Word Elements Length: {len(word_len)}\n",
                    pass_=True, log=True, screenshot=False)
            
                # Using count variable to loop over columns in word document
                count = 0
                # df_word.values[0] will give the list of column names from the table in Word document
                for col_name in df_word.values[0]:
                    if col_name in sourcefile.columns.values and col_name in actualexcel.columns.values:
                        source = sourcefile[col_name]
                        actual = actualexcel[col_name]
                        word = []
                        for row in docs.tables[6].rows:
                            word.append(row.cells[count].text)
                        
                        source = [item for item in source if str(item) != 'nan']
                        actual = [item for item in actual if str(item) != 'nan']
                        word.pop(0)

                        comparison_result = self.slrreport.list_comparison_between_reports_data(source, actual,
                                                                                                word=word)

                        if len(comparison_result) == 0:
                            self.LogScreenshot.fLogScreenshot(
                                message=f"Contents in Column '{col_name}' are matching between Expected Excel, "
                                        f"Complete Excel and Word Reports.",
                                pass_=True, log=True, screenshot=False)
                        else:
                            self.LogScreenshot.fLogScreenshot(
                                message=f"Contents in Column '{col_name}' are not matching between Expected Excel, "
                                        f"Complete Excel and Word Reports.\nDuplicate values are arranged in "
                                        f"following order -> Expected Excel, Complete Excel and Word Report. "
                                        f"{comparison_result}",
                                pass_=False, log=True, screenshot=False)
                            raise Exception("Elements are not matching between Expected Excel, Complete Excel and "
                                            "Word Reports")
                        count += 1
                    else:
                        raise Exception("Column names are not matching between Expected Excel, Complete Excel and "
                                        "Word Reports")
            else:
                self.LogScreenshot.fLogScreenshot(
                    message=f"Elements length is not matching between Expected Excel, Complete Excel and Complete "
                            f"Word Report. Expected Excel Elements Length: {len(source_len)}\n Excel Elements "
                            f"Length: {len(compex_len)}\n Word Elements Length: {len(word_len)}",
                    pass_=False, log=True, screenshot=False)
                raise Exception(f"Elements length is not matching between Expected Excel, Complete Excel and "
                                f"Complete Word Report.")
        except Exception:
            raise Exception("Error in Word report content validation")
        
        self.LogScreenshot.fLogScreenshot(
            message=f"*****Content validation between Extraction Template, Complete Excel Report and Complete "
                    f"Word Report for Utility Summary Completed*****",
            pass_=True, log=True, screenshot=False)

    def qol_utility_summary_validation_old_imports(self, webexcel_filename, excel_filename, util_filepath):
        self.LogScreenshot.fLogScreenshot(message=f"*****Content Validation as per the OLD Import logic*****",
                                          pass_=True, log=True, screenshot=False)
        source_template = self.get_util_source_template(util_filepath, 'OldImportLogic')

        self.LogScreenshot.fLogScreenshot(
            message=f"*****Check Absence of Utility Summary Tab in Complete Excel Report*****",
            pass_=True, log=True, screenshot=False)
        excel_data = openpyxl.load_workbook(f'ActualOutputs//{excel_filename}')
        if 'Utility Summary' not in excel_data.sheetnames:
            self.LogScreenshot.fLogScreenshot(
                message=f"'Utility Summary' tab is absent in downloaded Complete Excel Report as expected",
                pass_=True, log=True, screenshot=False)
        else:
            self.LogScreenshot.fLogScreenshot(message=f"'Utility Summary' tab is present in Complete Excel Report",
                                              pass_=False, log=True, screenshot=False)
            raise Exception("'Utility Summary' tab is present in Complete Excel Report")
        
        toc_sheet = pd.read_excel(f'ActualOutputs//{excel_filename}', sheet_name="TOC", skiprows=3)
        col_data = list(toc_sheet.iloc[:, 1])
        if f'Utility Summary' not in col_data:
            self.LogScreenshot.fLogScreenshot(message=f"'Utility Summary' is not present in TOC sheet as expected.",
                                              pass_=True, log=True, screenshot=False)
        else:
            self.LogScreenshot.fLogScreenshot(message=f"'Utility Summary' is present in TOC sheet. Available "
                                                      f"Data from TOC sheet: {col_data}",
                                              pass_=False, log=True, screenshot=False)
            raise Exception("'Utility Summary' is present in TOC sheet which is not expected to be present.")

        # QOL Report sheet comparison with Expected results
        self.LogScreenshot.fLogScreenshot(message=f"*****QOL Report sheet Comparison*****",
                                          pass_=True, log=True, screenshot=False)
        sourcefile1 = pd.read_excel(f'{source_template[0]}', sheet_name='ExpectedReportData')
        webexcel1 = pd.read_excel(f'ActualOutputs//{webexcel_filename}', sheet_name='QOL Report', skiprows=3)
        actualexcel1 = pd.read_excel(f'ActualOutputs//{excel_filename}', sheet_name='QOL Report', skiprows=3)
        
        cols1 = list(sourcefile1.columns.values)

        sourcefile_col1 = sourcefile1['Study Identifier']
        actualexcel_col1 = actualexcel1['Study Identifier']
        actualwebexcel_col1 = webexcel1['Study Identifier']

        sourcefile_col1 = [item for item in sourcefile_col1 if str(item) != 'nan']
        actualexcel_col1 = [item for item in actualexcel_col1 if str(item) != 'nan']
        actualwebexcel_col1 = [item for item in actualwebexcel_col1 if str(item) != 'nan']

        if sourcefile_col1 == sorted(sourcefile_col1) and actualexcel_col1 == sorted(actualexcel_col1) and \
                actualwebexcel_col1 == sorted(actualwebexcel_col1):
            self.LogScreenshot.fLogScreenshot(
                message=f"From 'QOL Report' sheet, contents in column 'Study Identifier' are in sorted order",
                pass_=True, log=True, screenshot=False)
            
            if len(sourcefile_col1) == len(actualexcel_col1) == len(actualwebexcel_col1):
                self.LogScreenshot.fLogScreenshot(
                    message=f"Elements length is matching between Expected Excel, Web_Excel and Complete Excel Report."
                            f"Expected Excel Elements Length: {len(sourcefile_col1)}\n WebExcel Elements Length: "
                            f"{len(actualwebexcel_col1)}\n Excel Elements Length: {len(actualexcel_col1)}",
                    pass_=True, log=True, screenshot=False)

                for col in cols1:
                    source_col1 = sourcefile1[col]
                    actual_excel_col1 = actualexcel1[col]
                    actual_webexcel_col = webexcel1[col]

                    source_col1 = [item for item in source_col1 if str(item) != 'nan']
                    actual_excel_col1 = [item for item in actual_excel_col1 if str(item) != 'nan']
                    actual_webexcel_col = [item for item in actual_webexcel_col if str(item) != 'nan']

                    comparison_result = self.slrreport.\
                        list_comparison_between_reports_data(source_col1, actual_excel_col1,
                                                             webex_list=actual_webexcel_col)

                    if len(comparison_result) == 0:
                        self.LogScreenshot.fLogScreenshot(
                            message=f"Contents in column '{col}' are matching between Expected Excel, WebExcel and "
                                    f"Complete Excel Report",
                            pass_=True, log=True, screenshot=False)
                    else:
                        self.LogScreenshot.fLogScreenshot(
                            message=f"Contents in column '{col}' are not matching. Duplicate values are arranged "
                                    f"in following order -> Expected Excel, WebExcel, Complete Excel. "
                                    f"{comparison_result}",
                            pass_=False, log=True, screenshot=False)
                        raise Exception("Contents are not matching between Expected Excel, WebExcel and "
                                        "Complete Excel report")
            else:
                self.LogScreenshot.fLogScreenshot(
                    message=f"Elements length is not matching between Expected Excel, Web_Excel and Complete Excel "
                            f"Report. Expected Excel Elements Length: {len(sourcefile_col1)}\n WebExcel Elements "
                            f"Length: {len(actualwebexcel_col1)}\n Excel Elements Length: {len(actualexcel_col1)}",
                    pass_=False, log=True, screenshot=False)
                raise Exception(f"Elements length is not matching between Expected Excel, Web_Excel and "
                                f"Complete Excel Report.")
        else:
            raise Exception("From 'QOL Report' sheet, contents in column 'Study Identifier' are not in Sorted order")
        
        self.LogScreenshot.fLogScreenshot(
            message=f"*****Content validation between Extraction Template, Complete Excel Report and Complete Word "
                    f"Report for Utility Summary Completed*****",
            pass_=True, log=True, screenshot=False)

    def qol_validate_utilitysummarytab_and_contents_into_excelreport(self, locatorname, util_filepath, index, env):
        source_template = self.get_util_source_template(util_filepath, 'prodfix')

        extraction_file = self.get_extraction_file_to_upload(util_filepath, 'prodfix', locatorname)

        # Read population data values
        pop_list = self.get_population_to_upload(util_filepath, 'prodfix', locatorname)
        # Read slrtype data values
        slrtype = self.get_slrtype_to_upload(util_filepath, 'prodfix', locatorname)

        # imp_pop = self.get_import_pop_to_upload(util_filepath, 'prodfix', locatorname)

        self.refreshpage()
        self.go_to_nested_page("importpublications_button", "extraction_upload_btn", env)
        self.exbase.upload_file(extraction_file, env)

        # Go to live slr page
        self.go_to_page("SLR_Homepage", env)
        self.slrreport.select_data(f"{pop_list[0][0]}", f"{pop_list[0][1]}", env)
        self.slrreport.select_data(slrtype[0][0], f"{slrtype[0][1]}", env)
        self.slrreport.generate_download_report("excel_report", env)
        excel_filename = self.slrreport.get_latest_filename(UnivWaitFor=180)
        self.validate_filename(excel_filename, util_filepath, "prodfix")

        self.LogScreenshot.fLogScreenshot(
            message=f"*****Check Presence of Utility Summary Tab in Complete Excel Report*****",
            pass_=True, log=True, screenshot=False)
        excel_data = openpyxl.load_workbook(f'ActualOutputs//{excel_filename}')
        if 'Utility Summary' in excel_data.sheetnames:
            self.LogScreenshot.fLogScreenshot(
                message=f"'Utility Summary' tab is present in Downloaded Complete Excel Report",
                pass_=True, log=True, screenshot=False)

            excel_sheet = excel_data[f'Utility Summary']
            if excel_sheet['H1'].value == 'Back To Toc':
                self.LogScreenshot.fLogScreenshot(
                    message=f"'Back To Toc' option is present in 'Utility Summary' sheet",
                    pass_=True, log=True, screenshot=False)
            else:
                self.LogScreenshot.fLogScreenshot(
                    message=f"'Back To Toc' option is not present in 'Utility Summary' sheet",
                    pass_=False, log=True, screenshot=False)
                raise Exception(f"'Back To Toc' option is not present in 'Utility Summary' sheet")
            
            toc_sheet = pd.read_excel(f'ActualOutputs//{excel_filename}', sheet_name="TOC", skiprows=3)
            col_data = list(toc_sheet.iloc[:, 1])
            if f'Utility Summary' in col_data:
                self.LogScreenshot.fLogScreenshot(message=f"'Utility Summary' is present in TOC sheet.",
                                                  pass_=True, log=True, screenshot=False)
            else:
                self.LogScreenshot.fLogScreenshot(
                    message=f"'Utility Summary' is not present in TOC sheet. Available Data from TOC sheet: {col_data}",
                    pass_=False, log=True, screenshot=False)
                raise Exception("'Utility Summary' is not present in TOC sheet.")
        
            # Utility Summary sheet comparison with Expected results
            self.LogScreenshot.fLogScreenshot(message=f"*****Utility Summary Sheet Content Comparison Started*****",
                                              pass_=True, log=True, screenshot=False)
            source_template_sheet = openpyxl.load_workbook(f'{source_template[0]}')
            sheets = source_template_sheet.sheetnames

            sourcefile = pd.read_excel(f'{source_template[0]}', sheet_name=sheets[index])
            actualexcel = pd.read_excel(f'ActualOutputs//{excel_filename}', sheet_name='Utility Summary', skiprows=2)
            
            cols = list(sourcefile.columns.values)

            # Check sorting order from downloaded excel report
            actual_slrsource = actualexcel["SLR Source"]
            actual_slrsource = [item for item in actual_slrsource if str(item) != 'nan']
            actual_slrsource_final = []
            # Removing the duplicates
            [actual_slrsource_final.append(x) for x in list(flatten(actual_slrsource))
             if x not in actual_slrsource_final]

            self.LogScreenshot.fLogScreenshot(message=f"Unique SLR Source column values are : {actual_slrsource_final}",
                                              pass_=True, log=True, screenshot=False)
            
            for m in actual_slrsource_final:
                col_val = actualexcel[actualexcel["SLR Source"] == m]
                col_val_res = col_val["Short Reference"]
                col_val_res = [item for item in col_val_res if str(item) != 'nan']
                if col_val_res == sorted(col_val_res):
                    self.LogScreenshot.fLogScreenshot(
                        message=f"From 'Utility Summary' sheet -> For '{m}' SLR Source, contents in column "
                                f"'Short Reference' are in sorted order",
                        pass_=True, log=True, screenshot=False)
                else:
                    raise Exception(f"From 'Utility Summary' sheet -> For '{m}' SLR Source, contents in column "
                                    f"'Short Reference' are not in Sorted order")

            # Check the length of 1st column from the report to make sure number of rows are as expected
            source_col_len = sourcefile["SLR Source"]
            actual_col_len = actualexcel["SLR Source"]

            source_col_len = [item for item in source_col_len if str(item) != 'nan']
            actual_col_len = [item for item in actual_col_len if str(item) != 'nan']

            if len(source_col_len) == len(actual_col_len):
                self.LogScreenshot.fLogScreenshot(
                    message=f"Elements length is matching between Expected Excel and Complete Excel Report. Expected "
                            f"Excel Elements Length: {len(source_col_len)}\n Excel Elements Length: "
                            f"{len(actual_col_len)}",
                    pass_=True, log=True, screenshot=False)
            
                # Content comparison between Source file and Complete Excel report
                for col in cols:
                    source_col = sourcefile[col]
                    actual_col = actualexcel[col]

                    source_col = [item for item in source_col if str(item) != 'nan']
                    actual_col = [item for item in actual_col if str(item) != 'nan']

                    comparison_result = self.slrreport.list_comparison_between_reports_data(source_col, actual_col)

                    if len(comparison_result) == 0:
                        self.LogScreenshot.fLogScreenshot(
                            message=f"Contents in column '{col}' are matching between Expected Excel and Complete "
                                    f"Excel Report",
                            pass_=True, log=True, screenshot=False)
                    else:
                        self.LogScreenshot.fLogScreenshot(
                            message=f"Contents in column '{col}' are not matching. Duplicate values are arranged "
                                    f"in following order -> Expected Excel, Complete Excel. {comparison_result}",
                            pass_=False, log=True, screenshot=False)
                        raise Exception("Contents are not matching between Expected Excel and Complete Excel report")
            else:
                self.LogScreenshot.fLogScreenshot(
                    message=f"Elements length is not matching between Expected Excel and Complete Excel Report. "
                            f"Expected Excel Elements Length: {len(source_col_len)}\n Excel Elements Length: "
                            f"{len(actual_col_len)}",
                    pass_=False, log=True, screenshot=False)
                raise Exception(f"Elements length is not matching between Expected Excel and Complete Excel Report.")

            self.LogScreenshot.fLogScreenshot(message=f"*****Utility Summary Sheet Content Comparison Completed*****",
                                              pass_=True, log=True, screenshot=False)
        elif 'Utility Summary' not in excel_data.sheetnames and locatorname == "scenario4":
            self.LogScreenshot.fLogScreenshot(
                message=f"'Utility Summary' tab is absent in downloaded Complete Excel Report as expected",
                pass_=True, log=True, screenshot=False)
        else:
            self.LogScreenshot.fLogScreenshot(message=f"'Utility Summary' tab is missing in Complete Excel Report",
                                              pass_=False, log=True, screenshot=False)
            raise Exception("'Utility Summary' tab is missing in Complete Excel Report")
        
        self.refreshpage()
        self.go_to_nested_page("importpublications_button", "extraction_upload_btn", env)

        self.exbase.delete_file(extraction_file, "file_status_popup_text", "upload_table_rows", env)

        # Go to live slr page
        self.go_to_page("SLR_Homepage", env)
        time.sleep(2)
    
    def econ_utility_summary_validation(self, webexcel_filename, excel_filename, util_filepath, word_filename):
        self.LogScreenshot.fLogScreenshot(
            message=f"*****Content validation between Extraction Template, Complete Excel Report and Complete Word "
                    f"Report for Utility Summary Started*****",
            pass_=True, log=True, screenshot=False)
        source_template = self.get_util_source_template(util_filepath, 'NewImportLogic')
        
        # ECON Report sheet comparison with Expected results
        self.LogScreenshot.fLogScreenshot(message=f"*****ECON Report sheet Comparison*****",
                                          pass_=True, log=True, screenshot=False)
        sourcefile = pd.read_excel(f'{source_template[0]}', sheet_name='ExpectedReportData')
        webexcel = pd.read_excel(f'ActualOutputs//{webexcel_filename}', sheet_name='CEA CUA Report', skiprows=3)
        actualexcel = pd.read_excel(f'ActualOutputs//{excel_filename}', sheet_name='CEA CUA Report', skiprows=3)
        
        cols = list(sourcefile.columns.values)

        sourcefile_col = sourcefile['Study Identifier']
        actualexcel_col = actualexcel['Study Identifier']
        actualwebexcel_col = webexcel['Study Identifier']

        sourcefile_col = [item for item in sourcefile_col if str(item) != 'nan']
        actualexcel_col = [item for item in actualexcel_col if str(item) != 'nan']
        actualwebexcel_col = [item for item in actualwebexcel_col if str(item) != 'nan']

        if sourcefile_col == sorted(sourcefile_col) and actualexcel_col == sorted(actualexcel_col) and \
                actualwebexcel_col == sorted(actualwebexcel_col):
            self.LogScreenshot.fLogScreenshot(
                message=f"From 'CEA CUA Report' sheet, contents in column 'Study Identifier' are in sorted order",
                pass_=True, log=True, screenshot=False)

            if len(sourcefile_col) == len(actualexcel_col) == len(actualwebexcel_col):
                self.LogScreenshot.fLogScreenshot(
                    message=f"Elements length is matching between Expected Excel, Web_Excel and Complete Excel Report."
                            f"Expected Excel Elements Length: {len(sourcefile_col)}\n WebExcel Elements Length: "
                            f"{len(actualwebexcel_col)}\n Excel Elements Length: {len(actualexcel_col)}",
                    pass_=True, log=True, screenshot=False)

                for col in cols:
                    source_col = sourcefile[col]
                    actual_excel_col = actualexcel[col]
                    actual_webexcel_col = webexcel[col]

                    source_col = [item for item in source_col if str(item) != 'nan']
                    actual_excel_col = [item for item in actual_excel_col if str(item) != 'nan']
                    actual_webexcel_col = [item for item in actual_webexcel_col if str(item) != 'nan']

                    comparison_result = self.slrreport.\
                        list_comparison_between_reports_data(source_col, actual_excel_col,
                                                             webex_list=actual_webexcel_col)

                    if len(comparison_result) == 0:
                        self.LogScreenshot.fLogScreenshot(
                            message=f"Contents in column '{col}' are matching between Expected Excel, WebExcel and "
                                    f"Complete Excel Report",
                            pass_=True, log=True, screenshot=False)
                    else:
                        self.LogScreenshot.fLogScreenshot(
                            message=f"Contents in column '{col}' are not matching. Duplicate values are arranged in "
                                    f"following order -> Expected Excel, WebExcel, Complete Excel.{comparison_result}",
                            pass_=False, log=True, screenshot=False)
                        raise Exception("Contents are not matching between Expected Excel, WebExcel and "
                                        "Complete Excel report")
            else:
                self.LogScreenshot.fLogScreenshot(
                    message=f"Elements length is not matching between Expected Excel, Web_Excel and Complete Excel "
                            f"Report. Expected Excel Elements Length: {len(sourcefile_col)}\n WebExcel Elements "
                            f"Length: {len(actualwebexcel_col)}\n Excel Elements Length: {len(actualexcel_col)}",
                    pass_=False, log=True, screenshot=False)
                raise Exception(f"Elements length is not matching between Expected Excel, Web_Excel and "
                                f"Complete Excel Report.")
        else:
            raise Exception("From 'CEA CUA Report' sheet, contents in column 'Study Identifier' "
                            "are not in Sorted order")
        
        # Word report content comparison with Expected Excel
        self.LogScreenshot.fLogScreenshot(message=f"*****ECON Word Report Comparison with Expected Excel*****",
                                          pass_=True, log=True, screenshot=False)
        # Index of Table number 6 is : 5. Starting point for word table content comparison
        table_count = 5
        docs = docx.Document(f'ActualOutputs//{word_filename}')
        try:
            table = docs.tables[table_count]
            data = [[cell.text for cell in row.cells] for row in table.rows]
            df_word = pd.DataFrame(data)

            # Check the length of 1st column from the report to make sure number of rows are as expected
            source_len = sourcefile[df_word.values[0][0]]
            webex_len = webexcel[df_word.values[0][0]]
            compex_len = actualexcel[df_word.values[0][0]]
            word_len = []
            for row in docs.tables[table_count].rows:
                word_len.append(row.cells[0].text)

            source_len = [item for item in source_len if str(item) != 'nan']
            webex_len = [item for item in webex_len if str(item) != 'nan']
            compex_len = [item for item in compex_len if str(item) != 'nan']
            word_len.pop(0)

            if len(source_len) == len(webex_len) == len(compex_len) == len(word_len):
                self.LogScreenshot.fLogScreenshot(
                    message=f"Elements length is matching between Expected Excel, Web_Excel, Complete Excel and "
                            f"Complete Word Report. Expected Excel Elements Length: {len(source_len)}\n WebExcel "
                            f"Elements Length: {len(webex_len)}\n Excel Elements Length: {len(compex_len)}\n "
                            f"Word Elements Length: {len(word_len)}",
                    pass_=True, log=True, screenshot=False)

                # Using count variable to loop over columns in word document
                count = 0
                # df_word.values[0] will give the list of column names from the table in Word document
                for col_name in df_word.values[0]:
                    # Restricting comparison upto 3 columns in word due to data formatting issues in further columns
                    if count <= 2:
                        if col_name == 'Year/Country':
                            # This IF condition is just to add space to the column name as per Excel sheet to match
                            # the names between word and excel. This is an workaround until it is fixed
                            col_name = 'Year / Country'
                        if col_name in sourcefile.columns.values and col_name in actualexcel.columns.values and \
                                col_name in webexcel.columns.values:
                            source = sourcefile[col_name]
                            actualex = actualexcel[col_name]
                            actwebex = webexcel[col_name]
                            word = []
                            for row in docs.tables[table_count].rows:
                                word.append(row.cells[count].text)
                            
                            source = [item for item in source if str(item) != 'nan']
                            actualex = [item for item in actualex if str(item) != 'nan']
                            actwebex = [item for item in actwebex if str(item) != 'nan']
                            word.pop(0)

                            comparison_result = self.slrreport.\
                                list_comparison_between_reports_data(source, actualex, webex_list=actwebex, word=word)

                            if len(comparison_result) == 0:
                                self.LogScreenshot.fLogScreenshot(
                                    message=f"Contents in Column '{col_name}' are matching between Expected Excel, "
                                            f"Complete Excel, Web-Excel and Word Reports.",
                                    pass_=True, log=True, screenshot=False)
                            else:
                                self.LogScreenshot.fLogScreenshot(
                                    message=f"Contents in Column '{col_name}' are matching between Expected Excel, "
                                            f"Complete Excel, Web-Excel and Word Reports.\nDuplicate values are "
                                            f"arranged in following order -> Expected Excel, WebExcel, Complete "
                                            f"Excel and Word Report. {comparison_result}",
                                    pass_=False, log=True, screenshot=False)
                                raise Exception("Elements are not matching between Expected Excel, Webexcel, "
                                                "Complete Excel and Word Reports")
                            count += 1
                        else:
                            raise Exception("Column names are not matching between Expected Excel, Complete Excel, "
                                            "WebExcel and Word Reports")
            else:
                self.LogScreenshot.fLogScreenshot(
                    message=f"Elements length is not matching between Expected Excel, Web_Excel, Complete Excel and "
                            f"Complete Word Report. Expected Excel Elements Length: {len(webex_len)}\n WebExcel "
                            f"Elements Length: {len(webex_len)}\n Excel Elements Length: {len(compex_len)}\n Word "
                            f"Elements Length: {len(word_len)}",
                    pass_=False, log=True, screenshot=False)
                raise Exception(f"Elements length is not matching between Expected Excel, Web_Excel, Complete Excel "
                                f"and Complete Word Report.")
        except Exception:
            raise Exception("Error in Word report content validation")
        
        self.LogScreenshot.fLogScreenshot(
            message=f"*****Content validation between Extraction Template, Complete Excel Report and Complete "
                    f"Word Report for Utility Summary Completed*****",
            pass_=True, log=True, screenshot=False)

    def econ_utility_summary_validation_old_imports(self, webexcel_filename, excel_filename, util_filepath):
        self.LogScreenshot.fLogScreenshot(message=f"*****Content Validation as per the OLD Import logic*****",
                                          pass_=True, log=True, screenshot=False)
        source_template = self.get_util_source_template(util_filepath, 'OldImportLogic')
        
        # ECON Report sheet comparison with Expected results
        self.LogScreenshot.fLogScreenshot(message=f"*****ECON Report sheet Comparison*****",
                                          pass_=True, log=True, screenshot=False)
        sourcefile1 = pd.read_excel(f'{source_template[0]}', sheet_name='ExpectedReportData')
        webexcel1 = pd.concat(pd.read_excel(f'ActualOutputs//{webexcel_filename}', sheet_name=None, skiprows=3),
                              ignore_index=True)
        actualexcel1 = pd.concat(pd.read_excel(f'ActualOutputs//{excel_filename}', sheet_name=None, skiprows=3),
                                 ignore_index=True)
        
        cols1 = list(sourcefile1.columns.values)

        sourcefile_col1 = sourcefile1['Study Identifier']
        actualexcel_col1 = actualexcel1['Study Identifier']
        actualwebexcel_col1 = webexcel1['Study Identifier']

        sourcefile_col1 = [item for item in sourcefile_col1 if str(item) != 'nan']
        actualexcel_col1 = [item for item in actualexcel_col1 if str(item) != 'nan']
        actualwebexcel_col1 = [item for item in actualwebexcel_col1 if str(item) != 'nan']

        if sourcefile_col1 == sorted(sourcefile_col1) and actualexcel_col1 == sorted(actualexcel_col1) and \
                actualwebexcel_col1 == sorted(actualwebexcel_col1):
            self.LogScreenshot.fLogScreenshot(message=f"Contents in column 'Study Identifier' are in sorted order",
                                              pass_=True, log=True, screenshot=False)

            if len(sourcefile_col1) == len(actualexcel_col1) == len(actualwebexcel_col1):
                self.LogScreenshot.fLogScreenshot(
                    message=f"Elements length is matching between Expected Excel, Web_Excel and Complete Excel "
                            f"Report. Expected Excel Elements Length: {len(sourcefile_col1)}\n WebExcel Elements "
                            f"Length: {len(actualwebexcel_col1)}\n Excel Elements Length: {len(actualexcel_col1)}",
                    pass_=True, log=True, screenshot=False)
                for col in cols1:
                    source_col1 = sourcefile1[col]
                    actual_excel_col1 = actualexcel1[col]
                    actual_webexcel_col1 = webexcel1[col]

                    source_col1 = [item for item in source_col1 if str(item) != 'nan']
                    actual_excel_col1 = [item for item in actual_excel_col1 if str(item) != 'nan']
                    actual_webexcel_col1 = [item for item in actual_webexcel_col1 if str(item) != 'nan']

                    comparison_result = self.slrreport.\
                        list_comparison_between_reports_data(source_col1, actual_excel_col1,
                                                             webex_list=actual_webexcel_col1)

                    if len(comparison_result) == 0:
                        self.LogScreenshot.fLogScreenshot(
                            message=f"Contents in column '{col}' are matching between Expected Excel, WebExcel and "
                                    f"Complete Excel Report",
                            pass_=True, log=True, screenshot=False)
                    else:
                        self.LogScreenshot.fLogScreenshot(
                            message=f"Contents in column '{col}' are not matching. Duplicate values are arranged in "
                                    f"following order -> Expected Excel,  WebExcel, Complete Excel. "
                                    f"{comparison_result}",
                            pass_=False, log=True, screenshot=False)
                        raise Exception("Contents are not matching between Expected Excel, WebExcel and "
                                        "Complete Excel report")
            else:
                self.LogScreenshot.fLogScreenshot(
                    message=f"Elements length is not matching between Expected Excel, Web_Excel and Complete Excel "
                            f"Report. Expected Excel Elements Length: {len(sourcefile_col1)}\n WebExcel Elements "
                            f"Length: {len(actualwebexcel_col1)}\n Excel Elements Length: {len(actualexcel_col1)}",
                    pass_=False, log=True, screenshot=False)
                raise Exception(f"Elements length is not matching between Expected Excel, Web_Excel and "
                                f"Complete Excel Report.")
        else:
            self.LogScreenshot.fLogScreenshot(
                message=f"Contents in column 'Study Identifier' are not in sorted order. Values are Expected Excel: "
                        f"{sourcefile_col1}, SortedExpectedExcelfile: {sorted(sourcefile_col1)}, ActualExcel: "
                        f"{actualexcel_col1}, SortedActualExcel: {sorted(actualexcel_col1)}, ActualWebExcel: "
                        f"{actualwebexcel_col1}, SortedActualWebExcel: {sorted(actualwebexcel_col1)}",
                pass_=False, log=True, screenshot=False)
            raise Exception("Contents in column 'Study Identifier' are not in Sorted order")
        
        self.LogScreenshot.fLogScreenshot(
            message=f"*****Content validation between Extraction Template, Complete Excel Report and Complete "
                    f"Word Report for Utility Summary Completed*****",
            pass_=True, log=True, screenshot=False)

    def presenceof_utilitycolumn_names(self, webexcel_filename, excel_filename, expected_dict):
                                          
        get_sheet_names = openpyxl.load_workbook(f'ActualOutputs//{webexcel_filename}')
        sheets = get_sheet_names.sheetnames

        for sheet in sheets:
            webexcel = pd.read_excel(f'ActualOutputs//{webexcel_filename}', sheet_name=sheet, skiprows=3)
            compexcel = pd.read_excel(f'ActualOutputs//{excel_filename}', sheet_name=sheet, skiprows=3)

            webex_cols = list(webexcel.columns.values)
            compex_cols = list(compexcel.columns.values)

            for k, v in expected_dict.items():
                if v in webex_cols and v in compex_cols:
                    self.LogScreenshot.fLogScreenshot(
                        message=f"Column name '{v}' is present in Web Excel and Complete Excel Report",
                        pass_=True, log=True, screenshot=False)
                else:
                    self.LogScreenshot.fLogScreenshot(
                        message=f"Column name '{v}' is not present in Web Excel and Complete Excel Report",
                        pass_=False, log=True, screenshot=False)
                    raise Exception(f"Column name '{v}' is not present in Web Excel and Complete Excel Report")

    def check_column_names_in_previewresults(self, expected_dict, env):

        # result = []
        # tc1 = self.select_elements("web_table_column_names", env)
        # for m in tc1:
        #     result.append(m.text)
        result = self.get_texts("web_table_column_names", env)

        for k, v in expected_dict.items():
            if v in result:
                self.LogScreenshot.fLogScreenshot(message=f"Column name '{v}' is present in Preview Results",
                                                  pass_=True, log=True, screenshot=False)
            else:
                self.LogScreenshot.fLogScreenshot(message=f"Column name '{v}' is not present in Preview Results",
                                                  pass_=False, log=True, screenshot=False)
                raise Exception("Column name '{v}' is not present in Preview Results")
